# ─── Word Document Export Functions ─────────────────────────────────
# Generates Word (.docx) reports, brief outlines, and trial binders.

import io
import json
import logging
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

logger = logging.getLogger(__name__)


def generate_word_report(state, case_name, module_filter=None):
    """
    Generates a Word document from the AgentState.
    Returns a BytesIO object.
    If module_filter is provided (set of module keys), only include those sections.
    Module keys: 'summary', 'strategy', 'witnesses', 'timeline', 'cross', 'direct',
                 'devils', 'evidence', 'consistency', 'elements', 'investigation',
                 'entities', 'voir_dire', 'research', 'medical'
    """
    # Cherry-pick helper: skip section if module_filter is set and this key is not in it
    _inc = lambda key: module_filter is None or key in module_filter

    doc = Document()

    # Title
    title = doc.add_heading(f"{case_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Case Summary
    if _inc('summary'):
        doc.add_heading("Case Summary", level=1)
        doc.add_paragraph(state.get("case_summary", "No summary available."))

    # 2. Defense Strategy
    if _inc('strategy'):
        doc.add_heading("Defense Strategy", level=1)
        doc.add_paragraph(state.get("strategy_notes", "No strategy notes available."))

    # 3. Witnesses
    doc.add_heading("Witness Analysis", level=1)
    witnesses = state.get("witnesses", [])
    if witnesses:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Goal'

        for w in witnesses:
            row_cells = table.add_row().cells
            row_cells[0].text = w.get("name", "Unknown")
            row_cells[1].text = w.get("type", "Unknown")
            row_cells[2].text = w.get("goal", "")
    else:
        doc.add_paragraph("No specific witnesses identified.")

    # 4. Timeline
    doc.add_heading("Chronological Timeline", level=1)
    timeline = state.get("timeline", [])
    if isinstance(timeline, str):
         doc.add_paragraph(timeline) # Fallback if raw string
    elif timeline:
         # Attempt to parse if it's a JSON string inside
         if isinstance(timeline, list):
             for event in timeline:
                 date = event.get("date", "")
                 time = event.get("time", "")
                 desc = event.get("event", "")
                 src = event.get("source", "")
                 doc.add_paragraph(f"{date} {time} - {desc} [Source: {src}]", style='List Bullet')
         else:
             doc.add_paragraph("Timeline data format error.")

    # 5. Cross Examination
    doc.add_heading("Cross-Examination Plan", level=1)
    cross = state.get("cross_examination_plan", [])
    if isinstance(cross, list):
        for item in cross:
            p = doc.add_paragraph(style='List Number')
            runner = p.add_run(f"Q: {item.get('question', '')}")
            runner.bold = True
            doc.add_paragraph(f"Rationale: {item.get('rationale', '')}")
            doc.add_paragraph(f"Source: {item.get('source', '')}")
            doc.add_paragraph("") # Spacer
    else:
        doc.add_paragraph(str(cross))

    # 6. Direct Examination
    doc.add_heading("Direct Examination Plan", level=1)
    direct = state.get("direct_examination_plan", [])
    if isinstance(direct, list):
        for item in direct:
            p = doc.add_paragraph(style='List Number')
            runner = p.add_run(f"Q: {item.get('question', '')}")
            runner.bold = True
            doc.add_paragraph(f"Goal: {item.get('goal', '')}")
            doc.add_paragraph("")
    else:
        doc.add_paragraph(str(direct))

    # 7. Devil's Advocate
    da = state.get("devils_advocate_notes", "")
    if da:
        doc.add_heading("Devil's Advocate Analysis", level=1)
        doc.add_paragraph(str(da))

    # 8. Evidence Foundations
    evidence = state.get("evidence_foundations", [])
    if evidence:
        doc.add_heading("Evidence Foundations", level=1)
        if isinstance(evidence, list):
            for e in evidence:
                if isinstance(e, dict):
                    p = doc.add_paragraph(style='List Bullet')
                    runner = p.add_run(str(e.get('item', e.get('evidence', ''))))
                    runner.bold = True
                    if e.get('foundation'):
                        doc.add_paragraph(f"Foundation: {e['foundation']}")
                    if e.get('admissibility'):
                        doc.add_paragraph(f"Admissibility: {e['admissibility']}")
                    if e.get('objections'):
                        doc.add_paragraph(f"Potential Objections: {e['objections']}")
                else:
                    doc.add_paragraph(str(e))
        else:
            doc.add_paragraph(str(evidence))

    # 9. Statement Conflicts
    conflicts = state.get("consistency_check", "")
    if conflicts:
        doc.add_heading("Statement Consistency & Contradictions", level=1)
        doc.add_paragraph(str(conflicts))

    # 10. Elements Map
    elements = state.get("elements_map", "")
    if elements:
        doc.add_heading("Legal Elements Map", level=1)
        doc.add_paragraph(str(elements))

    # 11. Investigation Plan
    investigation = state.get("investigation_plan", [])
    if investigation:
        doc.add_heading("Investigation Plan", level=1)
        if isinstance(investigation, list):
            for task in investigation:
                if isinstance(task, dict):
                    status = "[DONE]" if task.get('completed') else "[ ]"
                    doc.add_paragraph(f"{status} {task.get('task', task.get('description', str(task)))}", style='List Bullet')
                else:
                    doc.add_paragraph(f"[ ] {task}", style='List Bullet')
        else:
            doc.add_paragraph(str(investigation))

    # 12. Knowledge Graph Entities
    entities = state.get("entities", [])
    if entities:
        doc.add_heading("Knowledge Graph Entities", level=1)
        if isinstance(entities, list):
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Name'
            hdr[1].text = 'Type'
            hdr[2].text = 'Details'
            for ent in entities:
                if isinstance(ent, dict):
                    row = table.add_row().cells
                    row[0].text = str(ent.get('name', ''))
                    row[1].text = str(ent.get('type', ''))
                    row[2].text = str(ent.get('details', ent.get('description', '')))
        else:
            doc.add_paragraph(str(entities))

    # 13. Voir Dire Strategy
    voir_dire = state.get("voir_dire_strategy", "")
    if voir_dire:
        doc.add_heading("Voir Dire Strategy", level=1)
        doc.add_paragraph(str(voir_dire))

    # 14. Legal Research
    research = state.get("legal_research_data", [])
    research_summary = state.get("research_summary", "")
    if research_summary or research:
        doc.add_heading("Legal Research", level=1)
        if research_summary:
            doc.add_paragraph(str(research_summary))
        if isinstance(research, list):
            for r in research:
                if isinstance(r, dict):
                    doc.add_paragraph(f"{r.get('query', '')} -- {r.get('result', '')}", style='List Bullet')

    # 15. Medical Records Analysis
    medical = state.get("medical_records_analysis", {})
    if medical and isinstance(medical, dict):
        doc.add_heading("Medical Records Analysis", level=1)
        for key, value in medical.items():
            if value:
                doc.add_heading(key.replace('_', ' ').title(), level=2)
                doc.add_paragraph(str(value))

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_brief_outline(state, case_name):
    """
    Generates an IRAC-structured (Issue/Rule/Application/Conclusion)
    brief outline as a Word document. Returns a BytesIO object.
    """
    doc = Document()
    doc.add_heading(f"Legal Brief Outline -- {case_name}", level=0)
    doc.add_paragraph(f"Generated from CrimPro analysis")
    doc.add_paragraph("")

    # Strategy overview
    strategy = state.get("strategy_notes", "")
    if strategy:
        doc.add_heading("I. Case Strategy Overview", level=1)
        doc.add_paragraph(strategy if isinstance(strategy, str) else str(strategy))

    # IRAC for each charge
    charges = state.get("charges", [])
    elements = state.get("legal_elements", [])
    evidence = state.get("evidence_foundations", [])
    devils = state.get("devils_advocate", "")

    if charges and isinstance(charges, list):
        for i, charge in enumerate(charges):
            charge_name = charge if isinstance(charge, str) else charge.get("name", f"Charge {i+1}") if isinstance(charge, dict) else str(charge)
            doc.add_heading(f"II-{i+1}. {charge_name}", level=1)

            # ISSUE
            doc.add_heading("A. Issue", level=2)
            doc.add_paragraph(f"Whether sufficient evidence exists to sustain or defeat the charge of {charge_name}.")

            # RULE
            doc.add_heading("B. Rule", level=2)
            # Pull relevant elements
            relevant_elements = []
            if isinstance(elements, list):
                for el in elements:
                    if isinstance(el, dict):
                        el_name = el.get("element", el.get("name", ""))
                        el_desc = el.get("description", el.get("standard", ""))
                        el_strength = el.get("strength", "")
                        if el_name:
                            relevant_elements.append(f"* {el_name}: {el_desc} (Strength: {el_strength})")
            if relevant_elements:
                doc.add_paragraph("Required elements of the offense:")
                for el_text in relevant_elements:
                    doc.add_paragraph(el_text)
            else:
                doc.add_paragraph("[Insert applicable statutory elements and case law]")

            # APPLICATION
            doc.add_heading("C. Application", level=2)
            # Pull relevant evidence
            relevant_evidence = []
            if isinstance(evidence, list):
                for ev in evidence:
                    if isinstance(ev, dict):
                        ev_item = ev.get("evidence", ev.get("item", ""))
                        ev_foundation = ev.get("foundation", ev.get("basis", ""))
                        ev_admissible = ev.get("admissible", ev.get("admissibility", ""))
                        if ev_item:
                            relevant_evidence.append(f"* {ev_item}: {ev_foundation} (Admissibility: {ev_admissible})")
            if relevant_evidence:
                doc.add_paragraph("Key evidence applicable to this charge:")
                for ev_text in relevant_evidence:
                    doc.add_paragraph(ev_text)
            else:
                doc.add_paragraph("[Apply facts to legal elements]")

            # CONCLUSION
            doc.add_heading("D. Conclusion", level=2)
            doc.add_paragraph("[Insert conclusion for this charge based on analysis above]")
    else:
        doc.add_heading("II. Issues (No Charges Specified)", level=1)
        doc.add_paragraph("No specific charges were extracted. Add charges in the Summary tab.")

    # Devil's Advocate / Counter-arguments
    if devils:
        doc.add_heading("III. Anticipated Counter-Arguments", level=1)
        if isinstance(devils, str):
            doc.add_paragraph(devils)
        elif isinstance(devils, list):
            for item in devils:
                if isinstance(item, dict):
                    doc.add_paragraph(f"* {item.get('argument', item.get('point', str(item)))}")
                else:
                    doc.add_paragraph(f"* {str(item)}")

    # Witness Strategy
    witnesses = state.get("witnesses", [])
    if isinstance(witnesses, list) and witnesses:
        doc.add_heading("IV. Witness Strategy", level=1)
        for w in witnesses:
            if isinstance(w, dict):
                w_name = w.get("name", "Unknown")
                w_type = w.get("type", "")
                w_goal = w.get("goal", w.get("relevance", ""))
                doc.add_paragraph(f"* {w_name} ({w_type}): {w_goal}")

    # Summary / Case Theory
    summary = state.get("case_summary", "")
    if summary:
        doc.add_heading("V. Case Theory", level=1)
        doc.add_paragraph(summary if isinstance(summary, str) else str(summary))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =====================================================================
#  TRIAL BINDER GENERATOR
# =====================================================================

def generate_trial_binder(state, case_name, prep_type="trial", prep_name=""):
    """
    Generates a professional, courtroom-ready trial binder as a Word document.
    Organized into 13 tabbed sections mirroring a physical trial binder.
    Returns a BytesIO object.
    """
    doc = Document()

    # -- Helpers --------------------------------------------------------
    def safe(text):
        """Unicode-safe text for Word."""
        if not text:
            return ""
        t = str(text)
        t = t.replace('\u2014', '-').replace('\u2013', '-')
        t = t.replace('\u2018', "'").replace('\u2019', "'")
        t = t.replace('\u201c', '"').replace('\u201d', '"')
        t = t.replace('\u2022', '*').replace('\u2026', '...')
        return t

    def try_parse_json(content, fallback_type="list"):
        """Try to parse JSON from LLM output (may be wrapped in markdown)."""
        if not content:
            return [] if fallback_type == "list" else {}
        if isinstance(content, list):
            return content
        if isinstance(content, dict):
            return content
        text = str(content)
        try:
            # Strip markdown code fences
            text = re.sub(r'^```(?:json)?\s*', '', text, flags=re.MULTILINE)
            text = re.sub(r'```\s*$', '', text, flags=re.MULTILINE)
            text = text.strip()
            if fallback_type == "list":
                match = re.search(r'\[.*\]', text, re.DOTALL)
                if match:
                    return json.loads(match.group(0))
            else:
                match = re.search(r'\{.*\}', text, re.DOTALL)
                if match:
                    return json.loads(match.group(0))
        except (json.JSONDecodeError, Exception):
            pass
        return [] if fallback_type == "list" else {}

    def add_tab_header(doc, tab_num, title):
        """Creates a full-page tab divider separator sheet."""
        if tab_num > 1:
            doc.add_page_break()

        # Top spacing to vertically center content
        for _ in range(6):
            doc.add_paragraph()

        # Large tab number
        p_num = doc.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_num = p_num.add_run(f"TAB {tab_num}")
        run_num.bold = True
        run_num.font.size = Pt(48)
        run_num.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        # Horizontal rule
        p_rule = doc.add_paragraph()
        p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_rule = p_rule.add_run("\u2501" * 40)
        run_rule.font.size = Pt(14)
        run_rule.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

        # Section title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(title)
        run_title.bold = True
        run_title.font.size = Pt(24)
        run_title.font.color.rgb = RGBColor(0x34, 0x49, 0x5E)

        # Spacer
        doc.add_paragraph()
        doc.add_paragraph()

        # Instruction for physical binder assembly
        p_inst = doc.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_inst = p_inst.add_run("\u25b6  Insert physical tab divider here  \u25c0")
        run_inst.italic = True
        run_inst.font.size = Pt(11)
        run_inst.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Page break so content starts on the next page
        doc.add_page_break()

    def add_sub_heading(doc, text, level=1):
        h = doc.add_heading(safe(text), level=level)
        return h

    def add_body(doc, text, bold=False, italic=False):
        if not text:
            return
        p = doc.add_paragraph()
        run = p.add_run(safe(str(text)))
        run.bold = bold
        run.italic = italic
        return p

    def add_bullet(doc, text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        if bold_prefix:
            run = p.add_run(safe(bold_prefix))
            run.bold = True
            p.add_run(safe(str(text)))
        else:
            p.add_run(safe(str(text)))
        return p

    def shade_cells(row, color_hex):
        """Apply background shading to all cells in a row."""
        for cell in row.cells:
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): color_hex,
                qn('w:val'): 'clear'
            })
            shading.append(shd)

    def make_header_row(table, headers):
        """Bold and shade the header row of a table."""
        hdr = table.rows[0]
        for i, h_text in enumerate(headers):
            hdr.cells[i].text = safe(h_text)
            for paragraph in hdr.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        shade_cells(hdr, '2C3E50')
        # White text for dark header
        for cell in hdr.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def add_row(table, values, shade=False):
        row = table.add_row()
        for i, val in enumerate(values):
            row.cells[i].text = safe(str(val) if val else "")
            for paragraph in row.cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
        if shade:
            shade_cells(row, 'ECF0F1')
        return row

    def add_empty_note(doc, section_name):
        """Placeholder for sections with no data."""
        p = doc.add_paragraph()
        run = p.add_run(f"[No {section_name} data available -- run the corresponding analysis module to populate this section]")
        run.italic = True
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # -- Prep type labels -----------------------------------------------
    prep_type_labels = {
        "trial": "Trial Preparation",
        "prelim_hearing": "Preliminary Hearing",
        "motion_hearing": "Motion Hearing",
    }
    binder_label = prep_type_labels.get(prep_type, "Trial Preparation")

    # ==================================================================
    # TAB 1: COVER PAGE
    # ==================================================================
    for _ in range(4):
        doc.add_paragraph()  # Top margin spacing

    cover_title = doc.add_heading(safe(case_name), level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"\n{binder_label.upper()}")
    run.bold = True
    run.font.size = Pt(18)

    if prep_name:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(safe(prep_name))
        run2.font.size = Pt(14)
        run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()  # spacer

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Case type badge
    case_type = state.get("case_type", "criminal")
    case_type_display = {
        "criminal": "Criminal Defense",
        "criminal-juvenile": "Criminal - Juvenile",
        "civil-plaintiff": "Civil - Plaintiff",
        "civil-defendant": "Civil - Defendant",
        "civil-juvenile": "Civil - Juvenile",
    }.get(case_type, "Legal Matter")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"Case Type: {case_type_display}")
    run4.font.size = Pt(11)

    doc.add_paragraph()  # spacer
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("CONFIDENTIAL -- ATTORNEY WORK PRODUCT")
    run5.bold = True
    run5.font.size = Pt(10)
    run5.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # ==================================================================
    # TAB 2: TABLE OF CONTENTS
    # ==================================================================
    add_tab_header(doc, 2, "TABLE OF CONTENTS")

    toc_items = [
        ("Tab 1", "Cover Page"),
        ("Tab 2", "Table of Contents"),
        ("Tab 3", "Case Theory & Themes"),
        ("Tab 4", "Charges / Claims Checklist"),
        ("Tab 5", "Legal Elements Matrix"),
        ("Tab 6", "Witness Binder"),
        ("Tab 7", "Exhibit List & Foundations"),
        ("Tab 8", "Timeline of Events"),
        ("Tab 9", "Cross-Examination Outlines"),
        ("Tab 10", "Direct-Examination Outlines"),
        ("Tab 11", "Jury Selection (Voir Dire)"),
        ("Tab 12", "Risk Assessment & Devil's Advocate"),
        ("Tab 13", "Investigation Checklist"),
    ]

    toc_table = doc.add_table(rows=1, cols=3)
    toc_table.style = 'Table Grid'
    make_header_row(toc_table, ["Tab", "Section", "Status"])

    # Determine which sections have data
    section_data_keys = {
        "Tab 3": ["case_summary", "strategy_notes"],
        "Tab 4": ["charges"],
        "Tab 5": ["legal_elements"],
        "Tab 6": ["witnesses"],
        "Tab 7": ["evidence_foundations"],
        "Tab 8": ["timeline"],
        "Tab 9": ["cross_examination_plan"],
        "Tab 10": ["direct_examination_plan"],
        "Tab 11": ["voir_dire"],
        "Tab 12": ["devils_advocate_notes", "mock_jury_feedback", "consistency_check"],
        "Tab 13": ["investigation_plan"],
    }

    for i, (tab_label, section_title) in enumerate(toc_items):
        keys = section_data_keys.get(tab_label, [])
        has_data = any(state.get(k) for k in keys) if keys else True
        status = "Complete" if has_data else "Pending"
        status_icon = "[x]" if has_data else "[ ]"
        add_row(toc_table, [tab_label, section_title, f"{status_icon} {status}"], shade=(i % 2 == 0))

    # ==================================================================
    # TAB 3: CASE THEORY & THEMES
    # ==================================================================
    add_tab_header(doc, 3, "CASE THEORY & THEMES")

    case_summary = state.get("case_summary", "")
    strategy = state.get("strategy_notes", "")

    if case_summary:
        add_sub_heading(doc, "Case Summary", level=1)
        add_body(doc, case_summary)
    else:
        add_empty_note(doc, "Case Summary")

    if strategy:
        add_sub_heading(doc, "Defense Strategy & Key Themes", level=1)
        add_body(doc, strategy)
    else:
        add_empty_note(doc, "Strategy")

    # ==================================================================
    # TAB 4: CHARGES / CLAIMS CHECKLIST
    # ==================================================================
    add_tab_header(doc, 4, "CHARGES / CLAIMS CHECKLIST")

    charges = state.get("charges", [])
    if isinstance(charges, list) and charges:
        for ci, charge in enumerate(charges):
            if isinstance(charge, dict):
                charge_name = charge.get("name", f"Charge {ci + 1}")
                add_sub_heading(doc, f"Charge {ci + 1}: {charge_name}", level=2)

                # Details table
                detail_table = doc.add_table(rows=1, cols=2)
                detail_table.style = 'Table Grid'
                make_header_row(detail_table, ["Field", "Details"])

                fields = [
                    ("Offense", charge.get("name", "")),
                    ("Statute", charge.get("statute_number", charge.get("statute", ""))),
                    ("Severity", f"{charge.get('level', '')} {charge.get('class', '')}".strip()),
                    ("Statute Text", charge.get("statute_text", charge.get("text", ""))),
                    ("Jury Instructions", charge.get("jury_instructions", charge.get("instructions", ""))),
                ]
                for fi, (label, value) in enumerate(fields):
                    if value:
                        add_row(detail_table, [label, value], shade=(fi % 2 == 0))

                doc.add_paragraph()  # spacer
            elif isinstance(charge, str):
                add_bullet(doc, charge, bold_prefix=f"Charge {ci + 1}: ")
    else:
        add_empty_note(doc, "Charges / Claims")

    # ==================================================================
    # TAB 5: LEGAL ELEMENTS MATRIX
    # ==================================================================
    add_tab_header(doc, 5, "LEGAL ELEMENTS MATRIX")

    elements = state.get("legal_elements", [])
    elements_parsed = try_parse_json(elements, "list")

    if elements_parsed and isinstance(elements_parsed, list):
        el_table = doc.add_table(rows=1, cols=4)
        el_table.style = 'Table Grid'
        make_header_row(el_table, ["Charge", "Element", "Evidence", "Strength"])

        for ei, el in enumerate(elements_parsed):
            if isinstance(el, dict):
                add_row(el_table, [
                    el.get("charge", ""),
                    el.get("element", el.get("name", "")),
                    el.get("evidence", el.get("description", "")),
                    el.get("strength", el.get("status", "")),
                ], shade=(ei % 2 == 0))
    elif elements and isinstance(elements, str):
        # Fallback: raw text
        add_body(doc, elements)
    else:
        add_empty_note(doc, "Legal Elements")

    # ==================================================================
    # TAB 6: WITNESS BINDER
    # ==================================================================
    add_tab_header(doc, 6, "WITNESS BINDER")

    witnesses = state.get("witnesses", [])
    if isinstance(witnesses, list) and witnesses:
        # Master witness list table
        add_sub_heading(doc, "Master Witness List", level=1)
        w_table = doc.add_table(rows=1, cols=5)
        w_table.style = 'Table Grid'
        make_header_row(w_table, ["#", "Name", "Classification", "Goal / Role", "Contact"])

        for wi, w in enumerate(witnesses):
            if isinstance(w, dict):
                w_type = w.get("type", "Unknown")
                type_icon = {"State": "[HOSTILE]", "Defense": "[FRIENDLY]", "Swing": "[SWING]"}.get(w_type, "")
                add_row(w_table, [
                    str(wi + 1),
                    w.get("name", "Unknown"),
                    f"{w_type} {type_icon}",
                    w.get("goal", w.get("relevance", "")),
                    w.get("contact_info", "--"),
                ], shade=(wi % 2 == 0))

        # Per-witness detail sub-sections
        cross_plan = try_parse_json(state.get("cross_examination_plan", []), "list")
        direct_plan = try_parse_json(state.get("direct_examination_plan", []), "list")
        consistency = state.get("consistency_check", [])
        consistency_parsed = try_parse_json(consistency, "list")

        for wi, w in enumerate(witnesses):
            if not isinstance(w, dict):
                continue
            w_name = w.get("name", f"Witness {wi + 1}")
            w_type = w.get("type", "Unknown")

            doc.add_page_break()
            add_sub_heading(doc, f"Witness {wi + 1}: {w_name}", level=1)

            # Witness info box
            info_table = doc.add_table(rows=1, cols=2)
            info_table.style = 'Table Grid'
            make_header_row(info_table, ["Field", "Details"])
            add_row(info_table, ["Classification", w_type])
            add_row(info_table, ["Goal", w.get("goal", w.get("relevance", ""))], shade=True)
            add_row(info_table, ["Contact", w.get("contact_info", "Not available")])
            doc.add_paragraph()

            # Link matching cross-exam questions
            if w_type in ["State", "Swing"] and cross_plan:
                matched_cross = [c for c in cross_plan if isinstance(c, dict) and
                                 w_name.lower() in str(c.get("witness", "")).lower()]
                if matched_cross:
                    add_sub_heading(doc, f"Cross-Examination Outline -- {w_name}", level=2)
                    for cx in matched_cross:
                        topics = cx.get("topics", [])
                        if isinstance(topics, list):
                            for topic in topics:
                                if isinstance(topic, dict):
                                    add_sub_heading(doc, topic.get("title", "Topic"), level=3)
                                    questions = topic.get("questions", [])
                                    if isinstance(questions, list):
                                        for q in questions:
                                            if isinstance(q, dict):
                                                add_bullet(doc, f" -- {q.get('rationale', '')}",
                                                           bold_prefix=f"Q: {q.get('question', '')}")
                                                source = q.get("source", "")
                                                if source:
                                                    p = doc.add_paragraph()
                                                    run = p.add_run(f"    Source: {safe(source)}")
                                                    run.italic = True
                                                    run.font.size = Pt(8)
                                                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Link matching direct-exam questions
            if w_type in ["Defense", "Swing"] and direct_plan:
                matched_direct = [d for d in direct_plan if isinstance(d, dict) and
                                  w_name.lower() in str(d.get("witness", "")).lower()]
                if matched_direct:
                    add_sub_heading(doc, f"Direct-Examination Outline -- {w_name}", level=2)
                    for dx in matched_direct:
                        topics = dx.get("topics", [])
                        if isinstance(topics, list):
                            for topic in topics:
                                if isinstance(topic, dict):
                                    add_sub_heading(doc, topic.get("title", "Topic"), level=3)
                                    questions = topic.get("questions", [])
                                    if isinstance(questions, list):
                                        for q in questions:
                                            if isinstance(q, dict):
                                                add_bullet(doc, f" -- {q.get('goal', '')}",
                                                           bold_prefix=f"Q: {q.get('question', '')}")

            # Impeachment notes from consistency check
            if consistency_parsed and isinstance(consistency_parsed, list):
                related = [c for c in consistency_parsed if isinstance(c, dict) and
                           w_name.lower() in str(c).lower()]
                if related:
                    add_sub_heading(doc, f"Impeachment Notes -- {w_name}", level=2)
                    for conflict in related:
                        fact = conflict.get("fact", conflict.get("topic", ""))
                        detail = conflict.get("conflict", conflict.get("inconsistency", conflict.get("detail", "")))
                        add_bullet(doc, f": {detail}", bold_prefix=f"{fact}")
    else:
        add_empty_note(doc, "Witness")

    # ==================================================================
    # TAB 7: EXHIBIT LIST & FOUNDATIONS
    # ==================================================================
    add_tab_header(doc, 7, "EXHIBIT LIST & FOUNDATIONS")

    evidence = state.get("evidence_foundations", [])
    evidence_parsed = try_parse_json(evidence, "list")

    if evidence_parsed and isinstance(evidence_parsed, list):
        ex_table = doc.add_table(rows=1, cols=4)
        ex_table.style = 'Table Grid'
        make_header_row(ex_table, ["#", "Exhibit / Item", "Foundation (Admit)", "Attack (Suppress)"])

        for exi, ex in enumerate(evidence_parsed):
            if isinstance(ex, dict):
                add_row(ex_table, [
                    str(exi + 1),
                    ex.get("item", ex.get("evidence", ex.get("name", ""))),
                    ex.get("admissibility", ex.get("foundation", "")),
                    ex.get("attack", ex.get("objections", "")),
                ], shade=(exi % 2 == 0))
    elif evidence and isinstance(evidence, str):
        add_body(doc, evidence)
    else:
        add_empty_note(doc, "Evidence Foundations")

    # ==================================================================
    # TAB 8: TIMELINE OF EVENTS
    # ==================================================================
    add_tab_header(doc, 8, "TIMELINE OF EVENTS")

    timeline = state.get("timeline", [])
    timeline_parsed = try_parse_json(timeline, "list")

    if timeline_parsed and isinstance(timeline_parsed, list):
        tl_table = doc.add_table(rows=1, cols=4)
        tl_table.style = 'Table Grid'
        make_header_row(tl_table, ["Date", "Time", "Event", "Source"])

        for ti, evt in enumerate(timeline_parsed):
            if isinstance(evt, dict):
                add_row(tl_table, [
                    evt.get("date", ""),
                    evt.get("time", ""),
                    evt.get("event", evt.get("description", "")),
                    evt.get("source", ""),
                ], shade=(ti % 2 == 0))
    elif timeline and isinstance(timeline, str):
        add_body(doc, timeline)
    else:
        add_empty_note(doc, "Timeline")

    # ==================================================================
    # TAB 9: CROSS-EXAMINATION OUTLINES
    # ==================================================================
    add_tab_header(doc, 9, "CROSS-EXAMINATION OUTLINES")

    cross_data = state.get("cross_examination_plan", [])
    cross_parsed = try_parse_json(cross_data, "list")

    if cross_parsed and isinstance(cross_parsed, list):
        for cx in cross_parsed:
            if isinstance(cx, dict):
                witness = cx.get("witness", "Unknown Witness")
                add_sub_heading(doc, f"Cross: {witness}", level=1)

                topics = cx.get("topics", [])
                if isinstance(topics, list):
                    for topic in topics:
                        if isinstance(topic, dict):
                            add_sub_heading(doc, topic.get("title", "Topic"), level=2)
                            q_table = doc.add_table(rows=1, cols=3)
                            q_table.style = 'Table Grid'
                            make_header_row(q_table, ["Question (Leading)", "Source", "Rationale"])

                            questions = topic.get("questions", [])
                            if isinstance(questions, list):
                                for qi, q in enumerate(questions):
                                    if isinstance(q, dict):
                                        add_row(q_table, [
                                            q.get("question", ""),
                                            q.get("source", ""),
                                            q.get("rationale", ""),
                                        ], shade=(qi % 2 == 0))
                            doc.add_paragraph()
                else:
                    add_body(doc, str(topics))
            elif isinstance(cx, str):
                add_body(doc, cx)
    elif cross_data and isinstance(cross_data, str):
        # Fallback: raw text from LLM
        add_body(doc, cross_data)
    else:
        add_empty_note(doc, "Cross-Examination")

    # ==================================================================
    # TAB 10: DIRECT-EXAMINATION OUTLINES
    # ==================================================================
    add_tab_header(doc, 10, "DIRECT-EXAMINATION OUTLINES")

    direct_data = state.get("direct_examination_plan", [])
    direct_parsed = try_parse_json(direct_data, "list")

    if direct_parsed and isinstance(direct_parsed, list):
        for dx in direct_parsed:
            if isinstance(dx, dict):
                witness = dx.get("witness", "Unknown Witness")
                add_sub_heading(doc, f"Direct: {witness}", level=1)

                topics = dx.get("topics", [])
                if isinstance(topics, list):
                    for topic in topics:
                        if isinstance(topic, dict):
                            add_sub_heading(doc, topic.get("title", "Topic"), level=2)
                            q_table = doc.add_table(rows=1, cols=2)
                            q_table.style = 'Table Grid'
                            make_header_row(q_table, ["Question (Open-Ended)", "Goal"])

                            questions = topic.get("questions", [])
                            if isinstance(questions, list):
                                for qi, q in enumerate(questions):
                                    if isinstance(q, dict):
                                        add_row(q_table, [
                                            q.get("question", ""),
                                            q.get("goal", ""),
                                        ], shade=(qi % 2 == 0))
                            doc.add_paragraph()
                else:
                    add_body(doc, str(topics))
            elif isinstance(dx, str):
                add_body(doc, dx)
    elif direct_data and isinstance(direct_data, str):
        add_body(doc, direct_data)
    else:
        add_empty_note(doc, "Direct-Examination")

    # ==================================================================
    # TAB 11: JURY SELECTION (VOIR DIRE)
    # ==================================================================
    add_tab_header(doc, 11, "JURY SELECTION (VOIR DIRE)")

    voir_dire = state.get("voir_dire", {})
    vd_parsed = try_parse_json(voir_dire, "dict")

    if vd_parsed and isinstance(vd_parsed, dict):
        # Ideal juror
        ideal = vd_parsed.get("ideal_juror", "")
        if ideal:
            add_sub_heading(doc, "Ideal Juror Profile", level=1)
            add_body(doc, ideal)

        # Red flags
        red_flags = vd_parsed.get("red_flags", "")
        if red_flags:
            add_sub_heading(doc, "Red Flags -- Jurors to Strike", level=1)
            add_body(doc, red_flags)

        # Questions
        vd_questions = vd_parsed.get("questions", [])
        if isinstance(vd_questions, list) and vd_questions:
            add_sub_heading(doc, "Voir Dire Questions", level=1)
            vd_table = doc.add_table(rows=1, cols=3)
            vd_table.style = 'Table Grid'
            make_header_row(vd_table, ["#", "Question", "Goal"])

            for qi, q in enumerate(vd_questions):
                if isinstance(q, dict):
                    add_row(vd_table, [
                        str(qi + 1),
                        q.get("question", ""),
                        q.get("goal", ""),
                    ], shade=(qi % 2 == 0))
    elif voir_dire and isinstance(voir_dire, str):
        add_body(doc, voir_dire)
    else:
        add_empty_note(doc, "Voir Dire")

    # ==================================================================
    # TAB 12: RISK ASSESSMENT & DEVIL'S ADVOCATE
    # ==================================================================
    add_tab_header(doc, 12, "RISK ASSESSMENT & DEVIL'S ADVOCATE")

    # Devil's Advocate
    da = state.get("devils_advocate_notes", "")
    if da:
        add_sub_heading(doc, "Devil's Advocate Analysis", level=1)
        add_body(doc, da)
    else:
        add_empty_note(doc, "Devil's Advocate")

    # Mock Jury
    mock_jury = state.get("mock_jury_feedback", [])
    if isinstance(mock_jury, list) and mock_jury:
        add_sub_heading(doc, "Mock Jury Simulation Results", level=1)

        mj_table = doc.add_table(rows=1, cols=3)
        mj_table.style = 'Table Grid'
        make_header_row(mj_table, ["Juror Persona", "Verdict", "Reasoning"])

        for ji, juror in enumerate(mock_jury):
            if isinstance(juror, dict):
                verdict = juror.get("verdict", "")
                add_row(mj_table, [
                    juror.get("juror", f"Juror {ji + 1}"),
                    verdict,
                    juror.get("reaction", juror.get("reason", "")),
                ], shade=(ji % 2 == 0))

        # Tally
        guilty = sum(1 for j in mock_jury if isinstance(j, dict) and
                     "guilty" in str(j.get("verdict", "")).lower() and
                     "not" not in str(j.get("verdict", "")).lower())
        not_guilty = sum(1 for j in mock_jury if isinstance(j, dict) and
                         "not guilty" in str(j.get("verdict", "")).lower())
        undecided = len(mock_jury) - guilty - not_guilty

        p = doc.add_paragraph()
        run = p.add_run(f"\nVerdict Tally: Guilty {guilty} | Not Guilty {not_guilty} | Undecided {undecided}")
        run.bold = True

    # Consistency issues
    consistency = state.get("consistency_check", [])
    consistency_parsed_risk = try_parse_json(consistency, "list")

    if consistency_parsed_risk and isinstance(consistency_parsed_risk, list):
        add_sub_heading(doc, "Statement Inconsistencies & Contradictions", level=1)
        for ci, conflict in enumerate(consistency_parsed_risk):
            if isinstance(conflict, dict):
                fact = conflict.get("fact", conflict.get("topic", f"Issue {ci + 1}"))
                detail = conflict.get("conflict", conflict.get("inconsistency",
                         conflict.get("detail", str(conflict))))
                add_bullet(doc, f": {detail}", bold_prefix=safe(fact))
    elif consistency and isinstance(consistency, str):
        add_sub_heading(doc, "Statement Inconsistencies & Contradictions", level=1)
        add_body(doc, consistency)

    # ==================================================================
    # TAB 13: INVESTIGATION CHECKLIST
    # ==================================================================
    add_tab_header(doc, 13, "INVESTIGATION CHECKLIST")

    investigation = state.get("investigation_plan", [])
    if isinstance(investigation, list) and investigation:
        inv_table = doc.add_table(rows=1, cols=4)
        inv_table.style = 'Table Grid'
        make_header_row(inv_table, ["#", "Action Item", "Priority", "Status"])

        for ii, task in enumerate(investigation):
            if isinstance(task, dict):
                status = task.get("status", task.get("completed", ""))
                if isinstance(status, bool):
                    status = "Done" if status else "Open"
                # Check for user-edited flags
                if task.get("_user_completed"):
                    status = "Done (verified)"
                elif task.get("_user_added"):
                    status = "Open (manual)"
                add_row(inv_table, [
                    str(ii + 1),
                    task.get("action", task.get("task", task.get("description", ""))),
                    task.get("priority", ""),
                    status if status else "Open",
                ], shade=(ii % 2 == 0))
            elif isinstance(task, str):
                add_row(inv_table, [str(ii + 1), task, "", "Open"], shade=(ii % 2 == 0))
    elif investigation and isinstance(investigation, str):
        add_body(doc, investigation)
    else:
        add_empty_note(doc, "Investigation Plan")

    # -- Medical Records Appendix (if present) --------------------------
    medical = state.get("medical_records_analysis", {})
    if medical and isinstance(medical, dict) and len(medical) > 0:
        add_tab_header(doc, 14, "APPENDIX: MEDICAL RECORDS ANALYSIS")

        for key, value in medical.items():
            if value:
                add_sub_heading(doc, key.replace('_', ' ').title(), level=2)
                val_text = str(value)
                if len(val_text) > 8000:
                    val_text = val_text[:8000] + "\n\n[... truncated for binder length ...]"
                add_body(doc, val_text)

    # -- Legal Research Appendix (if present) ----------------------------
    research = state.get("legal_research_data", [])
    research_summary = state.get("research_summary", "")

    if research_summary or (isinstance(research, list) and research):
        tab_num = 15 if medical else 14
        add_tab_header(doc, tab_num, "APPENDIX: LEGAL RESEARCH")

        if research_summary:
            add_sub_heading(doc, "Research Summary", level=1)
            add_body(doc, research_summary)

        if isinstance(research, list):
            add_sub_heading(doc, "Research Results", level=1)
            for ri, r in enumerate(research):
                if isinstance(r, dict):
                    query = r.get("query", f"Query {ri + 1}")
                    result = r.get("result", "")
                    add_sub_heading(doc, f"Query: {query}", level=2)
                    result_text = str(result)
                    if len(result_text) > 5000:
                        result_text = result_text[:5000] + "\n\n[... truncated ...]"
                    add_body(doc, result_text)

    # -- Deposition Analysis Appendix (if present) ----------------------
    deposition = state.get("deposition_analysis", "")
    if deposition:
        tab_num = 14
        if medical:
            tab_num += 1
        if research_summary or (isinstance(research, list) and research):
            tab_num += 1
        add_tab_header(doc, tab_num, "APPENDIX: DEPOSITION ANALYSIS")
        add_body(doc, deposition)

    # -- Final page -----------------------------------------------------
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(6):
        p.add_run("\n")
    run = p.add_run("END OF TRIAL BINDER")
    run.bold = True
    run.font.size = Pt(16)
    p.add_run("\n\n")
    run2 = p.add_run(f"Generated by AllRise Beta on {datetime.now().strftime('%B %d, %Y')}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.add_run("\n")
    run3 = p.add_run("CONFIDENTIAL - ATTORNEY WORK PRODUCT")
    run3.bold = True
    run3.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
    run3.font.size = Pt(9)

    # -- Save -----------------------------------------------------------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
