# ─── Court Document Templates ──────────────────────────────────────
# Formats raw AI-drafted content into court-filing-ready Word documents.
# Supports multiple jurisdiction presets (Tennessee, Federal, Texas).

import io
import logging
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)


# Jurisdiction presets: margins, fonts, spacing, and formatting rules
JURISDICTION_PRESETS = {
    "tennessee_state": {
        "label": "Tennessee State Court",
        "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        "body_font": "Times New Roman",
        "body_size": 12,
        "heading_size": 14,
        "caption_size": 12,
        "line_spacing": 2.0,  # double-spaced
        "line_numbering": False,
        "page_size": "letter",
        "caption_style": "tn_state",  # Tennessee uses centered caption
        "certificate_template": (
            "CERTIFICATE OF SERVICE\n\n"
            "I hereby certify that a true and exact copy of the foregoing has been served upon "
            "all counsel of record by [METHOD OF SERVICE] on this the ___ day of __________, 20__.\n\n\n"
            "________________________________\n"
            "[ATTORNEY NAME]\n"
        ),
    },
    "federal": {
        "label": "Federal Court",
        "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        "body_font": "Times New Roman",
        "body_size": 12,
        "heading_size": 14,
        "caption_size": 12,
        "line_spacing": 2.0,
        "line_numbering": False,
        "page_size": "letter",
        "caption_style": "federal",
        "certificate_template": (
            "CERTIFICATE OF SERVICE\n\n"
            "I hereby certify that on this ___ day of __________, 20__, a true and correct copy "
            "of the foregoing document was served via the Court's CM/ECF system on all counsel of "
            "record who are registered CM/ECF users.\n\n\n"
            "________________________________\n"
            "[ATTORNEY NAME]\n"
        ),
    },
    "texas_state": {
        "label": "Texas State Court",
        "margins": {"top": 1.0, "bottom": 1.0, "left": 1.0, "right": 1.0},
        "body_font": "Times New Roman",
        "body_size": 12,
        "heading_size": 14,
        "caption_size": 12,
        "line_spacing": 2.0,
        "line_numbering": False,
        "page_size": "letter",
        "caption_style": "texas",
        "certificate_template": (
            "CERTIFICATE OF SERVICE\n\n"
            "I hereby certify that a true and correct copy of the above and foregoing document "
            "has been served on all attorneys of record in compliance with the Texas Rules of "
            "Civil Procedure on this the ___ day of __________, 20__.\n\n\n"
            "________________________________\n"
            "[ATTORNEY NAME]\n"
        ),
    },
}


def format_court_document(
    draft_content: str,
    document_title: str,
    jurisdiction: str = "tennessee_state",
    case_name: str = "",
    case_number: str = "",
    court_name: str = "",
    plaintiff: str = "",
    defendant: str = "",
    attorney_name: str = "",
    attorney_bar_number: str = "",
    attorney_firm: str = "",
    attorney_address: str = "",
    attorney_phone: str = "",
    attorney_email: str = "",
    include_certificate: bool = True,
    include_signature: bool = True,
    case_type: str = "civil",
) -> io.BytesIO:
    """
    Formats a raw draft into a court-filing-ready Word document.

    Args:
        draft_content: The raw text content from the AI draft generator
        document_title: e.g. "MOTION TO SUPPRESS EVIDENCE"
        jurisdiction: Key from JURISDICTION_PRESETS
        case_name: Full case name for caption
        case_number: Docket/case number
        court_name: Full court name
        plaintiff: Plaintiff name(s)
        defendant: Defendant name(s)
        attorney_*: Attorney information for signature block
        include_certificate: Whether to add Certificate of Service
        include_signature: Whether to add signature block
        case_type: "civil" or "criminal" (affects caption labels)

    Returns:
        BytesIO containing the formatted .docx file
    """
    preset = JURISDICTION_PRESETS.get(jurisdiction, JURISDICTION_PRESETS["tennessee_state"])

    doc = Document()

    # -- Page Setup -----------------------------------------------------
    for section in doc.sections:
        section.top_margin = Inches(preset["margins"]["top"])
        section.bottom_margin = Inches(preset["margins"]["bottom"])
        section.left_margin = Inches(preset["margins"]["left"])
        section.right_margin = Inches(preset["margins"]["right"])
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)

    # -- Style Setup ----------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = preset["body_font"]
    font.size = Pt(preset["body_size"])

    # Set default paragraph spacing to double
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    # Line spacing: 2.0 = double, 1.5 = one-and-a-half
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = preset["line_spacing"]

    # -- Helper Functions -----------------------------------------------
    def add_centered_text(text, size=12, bold=False, caps=False, spacing_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(spacing_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text.upper() if caps else text)
        run.font.name = preset["body_font"]
        run.font.size = Pt(size)
        run.bold = bold
        return p

    # -- Court Caption Block --------------------------------------------
    caption_style = preset.get("caption_style", "tn_state")

    # Court name header
    court_display = court_name or "[COURT NAME]"
    add_centered_text(court_display, size=preset["caption_size"], bold=True, caps=True, spacing_after=12)

    if caption_style == "tn_state":
        # Tennessee style: centered caption with vertical bar separator
        _build_tn_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)
    elif caption_style == "federal":
        _build_federal_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)
    elif caption_style == "texas":
        _build_texas_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)
    else:
        _build_tn_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)

    # -- Document Body --------------------------------------------------
    # Parse the draft content into paragraphs and apply court formatting
    _format_body_content(doc, draft_content, preset)

    # -- Signature Block ------------------------------------------------
    if include_signature:
        doc.add_paragraph()  # spacer
        _add_signature_block(
            doc, preset,
            attorney_name=attorney_name or "[ATTORNEY NAME]",
            bar_number=attorney_bar_number or "[BAR NUMBER]",
            firm=attorney_firm,
            address=attorney_address or "[ADDRESS]",
            phone=attorney_phone or "[PHONE]",
            email=attorney_email or "[EMAIL]",
        )

    # -- Certificate of Service -----------------------------------------
    if include_certificate:
        doc.add_page_break()
        _add_certificate_of_service(doc, preset, attorney_name or "[ATTORNEY NAME]")

    # -- Output ---------------------------------------------------------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _build_tn_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title):
    """Tennessee-style court caption with table layout."""
    # Use a table for the caption to get the vertical bar effect
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove table borders except the middle column
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.makeelement(qn('w:tcBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = borders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none',
                    qn('w:sz'): '0',
                })
                borders.append(border)
            tcPr.append(borders)

    # Left cell: parties
    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(2.75)

    # Plaintiff/State
    if case_type == "criminal":
        party_top = plaintiff or "STATE OF TENNESSEE"
        party_label_top = ""
        party_bottom = defendant or "[DEFENDANT NAME]"
        party_vs = "v."
    else:
        party_top = plaintiff or "[PLAINTIFF NAME]"
        party_label_top = "Plaintiff,"
        party_bottom = defendant or "[DEFENDANT NAME]"
        party_vs = "v."

    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(2)
    run = p1.add_run(party_top.upper())
    run.font.name = preset["body_font"]
    run.font.size = Pt(preset["caption_size"])
    run.bold = True

    if party_label_top:
        p1b = left_cell.add_paragraph()
        p1b.paragraph_format.line_spacing = 1.0
        p1b.paragraph_format.space_after = Pt(6)
        run1b = p1b.add_run(f"     {party_label_top}")
        run1b.font.name = preset["body_font"]
        run1b.font.size = Pt(preset["caption_size"])

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(party_vs)
    run2.font.name = preset["body_font"]
    run2.font.size = Pt(preset["caption_size"])

    p3 = left_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.line_spacing = 1.0
    run3 = p3.add_run(party_bottom.upper())
    run3.font.name = preset["body_font"]
    run3.font.size = Pt(preset["caption_size"])
    run3.bold = True

    if case_type != "criminal":
        p3b = left_cell.add_paragraph()
        p3b.paragraph_format.line_spacing = 1.0
        run3b = p3b.add_run("     Defendant.")
        run3b.font.name = preset["body_font"]
        run3b.font.size = Pt(preset["caption_size"])

    # Middle cell: just the bracket/line separator
    mid_cell = table.rows[0].cells[1]
    mid_cell.width = Inches(0.5)
    p_mid = mid_cell.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.line_spacing = 1.0
    run_mid = p_mid.add_run(")")
    run_mid.font.name = preset["body_font"]
    run_mid.font.size = Pt(preset["caption_size"])
    for _ in range(3):
        pm = mid_cell.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.line_spacing = 1.0
        pm.paragraph_format.space_after = Pt(2)
        rm = pm.add_run(")")
        rm.font.name = preset["body_font"]
        rm.font.size = Pt(preset["caption_size"])

    # Right cell: case number
    right_cell = table.rows[0].cells[2]
    right_cell.width = Inches(2.75)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_right.paragraph_format.line_spacing = 1.0
    p_right.paragraph_format.space_after = Pt(6)
    run_right = p_right.add_run(f"No. {case_number or '[CASE NUMBER]'}")
    run_right.font.name = preset["body_font"]
    run_right.font.size = Pt(preset["caption_size"])

    # Add spacing after table
    doc.add_paragraph()

    # Document title (centered, bold, caps)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    p_title.paragraph_format.line_spacing = 1.0
    run_title = p_title.add_run(document_title.upper())
    run_title.font.name = preset["body_font"]
    run_title.font.size = Pt(preset["heading_size"])
    run_title.bold = True
    run_title.underline = True


def _build_federal_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title):
    """Federal court caption style."""
    # Federal uses a similar table structure but with different labeling
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.makeelement(qn('w:tcBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = borders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none',
                    qn('w:sz'): '0',
                })
                borders.append(border)
            tcPr.append(borders)

    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(2.75)

    if case_type == "criminal":
        party_top = plaintiff or "UNITED STATES OF AMERICA"
    else:
        party_top = plaintiff or "[PLAINTIFF NAME]"

    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(2)
    run = p1.add_run(party_top.upper())
    run.font.name = preset["body_font"]
    run.font.size = Pt(preset["caption_size"])
    run.bold = True

    if case_type != "criminal":
        p1b = left_cell.add_paragraph()
        p1b.paragraph_format.line_spacing = 1.0
        p1b.paragraph_format.space_after = Pt(4)
        run1b = p1b.add_run("     Plaintiff,")
        run1b.font.name = preset["body_font"]
        run1b.font.size = Pt(preset["caption_size"])

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("v.")
    run2.font.name = preset["body_font"]
    run2.font.size = Pt(preset["caption_size"])

    p3 = left_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.line_spacing = 1.0
    run3 = p3.add_run((defendant or "[DEFENDANT NAME]").upper())
    run3.font.name = preset["body_font"]
    run3.font.size = Pt(preset["caption_size"])
    run3.bold = True

    if case_type != "criminal":
        p3b = left_cell.add_paragraph()
        p3b.paragraph_format.line_spacing = 1.0
        run3b = p3b.add_run("     Defendant.")
        run3b.font.name = preset["body_font"]
        run3b.font.size = Pt(preset["caption_size"])

    # Middle cell: bracket
    mid_cell = table.rows[0].cells[1]
    mid_cell.width = Inches(0.5)
    p_mid = mid_cell.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.line_spacing = 1.0
    run_mid = p_mid.add_run(")")
    run_mid.font.name = preset["body_font"]
    run_mid.font.size = Pt(preset["caption_size"])
    for _ in range(3):
        pm = mid_cell.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.line_spacing = 1.0
        pm.paragraph_format.space_after = Pt(2)
        rm = pm.add_run(")")
        rm.font.name = preset["body_font"]
        rm.font.size = Pt(preset["caption_size"])

    # Right cell: case details
    right_cell = table.rows[0].cells[2]
    right_cell.width = Inches(2.75)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_right.paragraph_format.line_spacing = 1.0
    p_right.paragraph_format.space_after = Pt(4)
    run_right = p_right.add_run(f"Case No. {case_number or '[CASE NUMBER]'}")
    run_right.font.name = preset["body_font"]
    run_right.font.size = Pt(preset["caption_size"])

    doc.add_paragraph()

    # Document title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    p_title.paragraph_format.line_spacing = 1.0
    run_title = p_title.add_run(document_title.upper())
    run_title.font.name = preset["body_font"]
    run_title.font.size = Pt(preset["heading_size"])
    run_title.bold = True
    run_title.underline = True


def _build_texas_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title):
    """Texas state court caption style."""
    # Texas uses a distinctive style with section symbols
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.makeelement(qn('w:tcBorders'), {})
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = borders.makeelement(qn(f'w:{border_name}'), {
                    qn('w:val'): 'none',
                    qn('w:sz'): '0',
                })
                borders.append(border)
            tcPr.append(borders)

    left_cell = table.rows[0].cells[0]
    left_cell.width = Inches(2.75)

    if case_type == "criminal":
        party_top = plaintiff or "THE STATE OF TEXAS"
    else:
        party_top = plaintiff or "[PLAINTIFF NAME]"

    p1 = left_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.line_spacing = 1.0
    p1.paragraph_format.space_after = Pt(4)
    run = p1.add_run(party_top.upper())
    run.font.name = preset["body_font"]
    run.font.size = Pt(preset["caption_size"])
    run.bold = True

    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 1.0
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run("v.")
    run2.font.name = preset["body_font"]
    run2.font.size = Pt(preset["caption_size"])

    p3 = left_cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.line_spacing = 1.0
    run3 = p3.add_run((defendant or "[DEFENDANT NAME]").upper())
    run3.font.name = preset["body_font"]
    run3.font.size = Pt(preset["caption_size"])
    run3.bold = True

    # Middle cell: section symbols (Texas signature)
    mid_cell = table.rows[0].cells[1]
    mid_cell.width = Inches(0.5)
    p_mid = mid_cell.paragraphs[0]
    p_mid.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mid.paragraph_format.line_spacing = 1.0
    run_mid = p_mid.add_run("\u00a7")
    run_mid.font.name = preset["body_font"]
    run_mid.font.size = Pt(preset["caption_size"])
    for _ in range(2):
        pm = mid_cell.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pm.paragraph_format.line_spacing = 1.0
        pm.paragraph_format.space_after = Pt(4)
        rm = pm.add_run("\u00a7")
        rm.font.name = preset["body_font"]
        rm.font.size = Pt(preset["caption_size"])

    # Right cell: court and case number
    right_cell = table.rows[0].cells[2]
    right_cell.width = Inches(2.75)
    p_right = right_cell.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_right.paragraph_format.line_spacing = 1.0
    p_right.paragraph_format.space_after = Pt(4)
    run_right = p_right.add_run(f"IN THE [___] COURT")
    run_right.font.name = preset["body_font"]
    run_right.font.size = Pt(preset["caption_size"])

    p_right2 = right_cell.add_paragraph()
    p_right2.paragraph_format.line_spacing = 1.0
    p_right2.paragraph_format.space_after = Pt(4)
    run_right2 = p_right2.add_run(f"Cause No. {case_number or '[CASE NUMBER]'}")
    run_right2.font.name = preset["body_font"]
    run_right2.font.size = Pt(preset["caption_size"])

    doc.add_paragraph()

    # Document title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    p_title.paragraph_format.line_spacing = 1.0
    run_title = p_title.add_run(document_title.upper())
    run_title.font.name = preset["body_font"]
    run_title.font.size = Pt(preset["heading_size"])
    run_title.bold = True
    run_title.underline = True


def _format_body_content(doc, content: str, preset: dict):
    """Parse raw draft content into properly formatted court document paragraphs."""
    if not content:
        return

    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if not line:
            # Empty line -- skip (double-spacing handles gaps)
            i += 1
            continue

        # Detect markdown-style headers -> section headings
        if line.startswith('### '):
            _add_court_heading(doc, line[4:], preset, level=3)
        elif line.startswith('## '):
            _add_court_heading(doc, line[3:], preset, level=2)
        elif line.startswith('# '):
            _add_court_heading(doc, line[2:], preset, level=1)
        # Roman numeral section headers (I., II., III., IV., etc.)
        elif re.match(r'^[IVXLCivxlc]+\.\s', line):
            _add_court_heading(doc, line, preset, level=2)
        # Numbered section headers (1., 2., etc. at start followed by uppercase)
        elif re.match(r'^\d+\.\s+[A-Z]', line) and len(line) < 100:
            _add_court_heading(doc, line, preset, level=2)
        # Bullet points
        elif line.startswith('- ') or line.startswith('* ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:])
            run.font.name = preset["body_font"]
            run.font.size = Pt(preset["body_size"])
        # Numbered list items
        elif re.match(r'^\d+[.)]\s', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+[.)]\s', '', line)
            run = p.add_run(text)
            run.font.name = preset["body_font"]
            run.font.size = Pt(preset["body_size"])
        # All-caps short lines -> treat as sub-headings
        elif line == line.upper() and len(line) < 80 and len(line) > 3:
            _add_court_heading(doc, line, preset, level=2)
        # Normal body paragraph
        else:
            p = doc.add_paragraph()
            # Check for bold markers **text**
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(part)
                run.font.name = preset["body_font"]
                run.font.size = Pt(preset["body_size"])
            # Set double spacing for body paragraphs
            p.paragraph_format.line_spacing = preset["line_spacing"]

        i += 1


def _add_court_heading(doc, text: str, preset: dict, level: int = 2):
    """Add a court-style heading (centered, bold, properly formatted)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level <= 2 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0

    # Clean markdown formatting
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    clean = re.sub(r'^#+\s*', '', clean)

    run = p.add_run(clean.upper() if level <= 2 else clean)
    run.font.name = preset["body_font"]
    run.font.size = Pt(preset["heading_size"] if level == 1 else preset["body_size"])
    run.bold = True
    if level <= 2:
        run.underline = True


def _add_signature_block(doc, preset, attorney_name, bar_number, firm, address, phone, email):
    """Add a professional attorney signature block."""
    p_respect = doc.add_paragraph()
    p_respect.paragraph_format.line_spacing = 1.0
    p_respect.paragraph_format.space_before = Pt(24)
    run_resp = p_respect.add_run("Respectfully submitted,")
    run_resp.font.name = preset["body_font"]
    run_resp.font.size = Pt(preset["body_size"])

    # Signature line
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.line_spacing = 1.0
    p_sig.paragraph_format.space_before = Pt(36)
    p_sig.paragraph_format.space_after = Pt(2)
    run_sig = p_sig.add_run("________________________________")
    run_sig.font.name = preset["body_font"]
    run_sig.font.size = Pt(preset["body_size"])

    # Attorney info lines
    info_lines = [attorney_name]
    if bar_number:
        info_lines.append(f"BPR #{bar_number}")
    if firm:
        info_lines.append(firm)
    if address:
        # Support multi-line address
        for addr_line in address.split('\n'):
            if addr_line.strip():
                info_lines.append(addr_line.strip())
    if phone:
        info_lines.append(f"Tel: {phone}")
    if email:
        info_lines.append(f"Email: {email}")

    for info_line in info_lines:
        p_info = doc.add_paragraph()
        p_info.paragraph_format.line_spacing = 1.0
        p_info.paragraph_format.space_after = Pt(0)
        p_info.paragraph_format.space_before = Pt(0)
        run_info = p_info.add_run(info_line)
        run_info.font.name = preset["body_font"]
        run_info.font.size = Pt(preset["body_size"])


def _add_certificate_of_service(doc, preset, attorney_name):
    """Add a Certificate of Service page."""
    template = preset.get("certificate_template", "")
    if not template:
        return

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    p_title.paragraph_format.line_spacing = 1.0
    run_title = p_title.add_run("CERTIFICATE OF SERVICE")
    run_title.font.name = preset["body_font"]
    run_title.font.size = Pt(preset["heading_size"])
    run_title.bold = True
    run_title.underline = True

    # Body of certificate -- replace placeholder with actual attorney name
    cert_body = template.replace("CERTIFICATE OF SERVICE\n\n", "")  # Remove the title we already added
    cert_body = cert_body.replace("[ATTORNEY NAME]", attorney_name)

    for line in cert_body.split('\n'):
        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = preset["line_spacing"]
            run = p.add_run(line)
            run.font.name = preset["body_font"]
            run.font.size = Pt(preset["body_size"])
        else:
            doc.add_paragraph()  # blank line


def generate_major_document_word(
    draft: dict,
    jurisdiction: str = "tennessee_state",
    attorney_info: dict = None,
    case_info: dict = None,
) -> io.BytesIO:
    """Generate a court-formatted Word document from a multi-section major document draft.

    Builds on the existing caption, body formatting, signature, and certificate
    infrastructure but handles multi-section documents with TOC and TOA.

    Args:
        draft: Full draft dict with outline, sections, citation_library, title.
        jurisdiction: Key from JURISDICTION_PRESETS.
        attorney_info: {name, bar_number, firm, address, phone, email}
        case_info: {plaintiff, defendant, case_number, court_name, case_type}

    Returns:
        BytesIO containing the formatted .docx file
    """
    attorney_info = attorney_info or {}
    case_info = case_info or {}
    preset = JURISDICTION_PRESETS.get(jurisdiction, JURISDICTION_PRESETS["tennessee_state"])

    doc = Document()

    # -- Page Setup --
    for section in doc.sections:
        section.top_margin = Inches(preset["margins"]["top"])
        section.bottom_margin = Inches(preset["margins"]["bottom"])
        section.left_margin = Inches(preset["margins"]["left"])
        section.right_margin = Inches(preset["margins"]["right"])
        section.page_height = Inches(11)
        section.page_width = Inches(8.5)

    # -- Style Setup --
    style = doc.styles['Normal']
    font = style.font
    font.name = preset["body_font"]
    font.size = Pt(preset["body_size"])
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = preset["line_spacing"]

    def _add_centered(text, size=12, bold=False, caps=False, spacing_after=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(spacing_after)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text.upper() if caps else text)
        run.font.name = preset["body_font"]
        run.font.size = Pt(size)
        run.bold = bold
        return p

    # -- Cover Page: Court Caption --
    document_title = draft.get("title", "DOCUMENT")
    court_name = case_info.get("court_name", "")
    plaintiff = case_info.get("plaintiff", "")
    defendant = case_info.get("defendant", "")
    case_number = case_info.get("case_number", "")
    case_type = case_info.get("case_type", "civil")

    _add_centered(
        court_name or "[COURT NAME]",
        size=preset["caption_size"], bold=True, caps=True, spacing_after=12,
    )

    caption_style = preset.get("caption_style", "tn_state")
    if caption_style == "federal":
        _build_federal_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)
    elif caption_style == "texas":
        _build_texas_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)
    else:
        _build_tn_caption(doc, preset, plaintiff, defendant, case_number, case_type, document_title)

    # -- Table of Contents --
    sections_data = draft.get("sections", [])
    outline = draft.get("outline", [])
    toc_source = sections_data if sections_data else outline

    if toc_source:
        doc.add_page_break()
        p_toc_title = doc.add_paragraph()
        p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_toc_title.paragraph_format.space_before = Pt(24)
        p_toc_title.paragraph_format.space_after = Pt(18)
        p_toc_title.paragraph_format.line_spacing = 1.0
        run_toc = p_toc_title.add_run("TABLE OF CONTENTS")
        run_toc.font.name = preset["body_font"]
        run_toc.font.size = Pt(preset["heading_size"])
        run_toc.bold = True
        run_toc.underline = True

        for sec in toc_source:
            p_entry = doc.add_paragraph()
            p_entry.paragraph_format.line_spacing = 1.0
            p_entry.paragraph_format.space_after = Pt(4)
            sec_num = sec.get("section_num", "")
            sec_title = sec.get("title", "")
            run_e = p_entry.add_run(f"{sec_num}. {sec_title}")
            run_e.font.name = preset["body_font"]
            run_e.font.size = Pt(preset["body_size"])

    # -- Table of Authorities --
    citation_library = draft.get("citation_library", [])
    if citation_library and sections_data:
        try:
            from core.nodes.major_docs import generate_table_of_authorities
            toa_text = generate_table_of_authorities(sections_data, citation_library)
        except Exception:
            toa_text = ""

        if toa_text and toa_text != "No authorities cited.":
            doc.add_page_break()
            for line in toa_text.split('\n'):
                if not line.strip():
                    doc.add_paragraph()
                elif line == "TABLE OF AUTHORITIES":
                    p_toa = doc.add_paragraph()
                    p_toa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_toa.paragraph_format.space_after = Pt(18)
                    p_toa.paragraph_format.line_spacing = 1.0
                    run_toa = p_toa.add_run(line)
                    run_toa.font.name = preset["body_font"]
                    run_toa.font.size = Pt(preset["heading_size"])
                    run_toa.bold = True
                    run_toa.underline = True
                elif line.rstrip().endswith(":") and not line.startswith("  "):
                    p_cat = doc.add_paragraph()
                    p_cat.paragraph_format.space_before = Pt(12)
                    p_cat.paragraph_format.space_after = Pt(4)
                    p_cat.paragraph_format.line_spacing = 1.0
                    run_cat = p_cat.add_run(line)
                    run_cat.font.name = preset["body_font"]
                    run_cat.font.size = Pt(preset["body_size"])
                    run_cat.bold = True
                else:
                    p_e = doc.add_paragraph()
                    p_e.paragraph_format.line_spacing = 1.0
                    p_e.paragraph_format.space_after = Pt(2)
                    run_e = p_e.add_run(line)
                    run_e.font.name = preset["body_font"]
                    run_e.font.size = Pt(preset["body_size"])

    # -- Body Sections --
    if sections_data:
        doc.add_page_break()
        for sec in sections_data:
            sec_num = sec.get("section_num", "")
            sec_title = sec.get("title", "")
            content = sec.get("content", "")

            _add_court_heading(doc, f"{sec_num}. {sec_title}", preset, level=2)
            if content:
                _format_body_content(doc, content, preset)

    # -- Signature Block --
    doc.add_paragraph()
    _add_signature_block(
        doc, preset,
        attorney_name=attorney_info.get("name") or "[ATTORNEY NAME]",
        bar_number=attorney_info.get("bar_number") or "[BAR NUMBER]",
        firm=attorney_info.get("firm", ""),
        address=attorney_info.get("address") or "[ADDRESS]",
        phone=attorney_info.get("phone") or "[PHONE]",
        email=attorney_info.get("email") or "[EMAIL]",
    )

    # -- Certificate of Service --
    doc.add_page_break()
    _add_certificate_of_service(
        doc, preset, attorney_info.get("name") or "[ATTORNEY NAME]",
    )

    # -- Output --
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_jurisdiction_list():
    """Return list of available jurisdiction presets for UI dropdown."""
    return [(key, preset["label"]) for key, preset in JURISDICTION_PRESETS.items()]
