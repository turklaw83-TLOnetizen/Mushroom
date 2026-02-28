# ─── Exhibit Index & Sticker Generator ─────────────────────────────
# Generates exhibit indices (Word) and printable exhibit stickers (PDF).

import io
import logging
import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

logger = logging.getLogger(__name__)


def generate_exhibit_index(case_files, case_name):
    """
    Generates a Word document with a numbered exhibit index table.
    case_files: list of file paths.
    Returns a BytesIO object.
    """
    doc = Document()

    # Title
    title = doc.add_heading(f"EXHIBIT INDEX -- {case_name}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Prepared: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph("")

    # Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Exhibit #'
    hdr[1].text = 'Description / File Name'
    hdr[2].text = 'Type'
    hdr[3].text = 'Date Added'

    # Bold header row
    for cell in table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for i, fpath in enumerate(case_files):
        fname = os.path.basename(fpath)
        ext = os.path.splitext(fname)[1].upper().replace('.', '')
        try:
            mtime = datetime.fromtimestamp(os.path.getmtime(fpath)).strftime('%Y-%m-%d')
        except Exception:
            mtime = ''

        row = table.add_row().cells
        row[0].text = f"Exhibit {i + 1}"
        row[1].text = fname
        row[2].text = ext
        row[3].text = mtime

    # Footer
    doc.add_paragraph("")
    doc.add_paragraph(f"Total Exhibits: {len(case_files)}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_exhibit_stickers(case_files, case_name):
    """
    Generates a PDF with generic printable exhibit labels.
    Each label shows the exhibit number, case name, and file name.
    Returns a BytesIO object.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Layout: 2 columns x 5 rows per page = 10 stickers per page
    sticker_w = 90  # mm
    sticker_h = 50  # mm
    margin_x = 15
    margin_y = 15
    gap_x = 5
    gap_y = 5

    for i, fpath in enumerate(case_files):
        col = i % 2
        row_on_page = (i // 2) % 5

        if i % 10 == 0:
            pdf.add_page()

        x = margin_x + col * (sticker_w + gap_x)
        y = margin_y + row_on_page * (sticker_h + gap_y)

        # Border
        pdf.rect(x, y, sticker_w, sticker_h)

        # Exhibit number -- large
        pdf.set_font('Helvetica', 'B', 18)
        pdf.set_xy(x + 3, y + 5)
        exhibit_label = f"EXHIBIT {i + 1}"
        try:
            pdf.cell(sticker_w - 6, 12, exhibit_label.encode('latin-1', 'replace').decode('latin-1'), 0, 1, 'C')
        except Exception:
            pdf.cell(sticker_w - 6, 12, f"EXHIBIT {i + 1}", 0, 1, 'C')

        # Horizontal line
        pdf.line(x + 5, y + 20, x + sticker_w - 5, y + 20)

        # Case name
        pdf.set_font('Helvetica', '', 9)
        pdf.set_xy(x + 3, y + 22)
        safe_case = case_name.encode('latin-1', 'replace').decode('latin-1')[:40]
        pdf.cell(sticker_w - 6, 6, safe_case, 0, 1, 'C')

        # File name
        fname = os.path.basename(fpath)
        pdf.set_font('Helvetica', 'I', 8)
        pdf.set_xy(x + 3, y + 30)
        safe_fname = fname.encode('latin-1', 'replace').decode('latin-1')[:45]
        pdf.cell(sticker_w - 6, 5, safe_fname, 0, 1, 'C')

        # Date
        pdf.set_font('Helvetica', '', 7)
        pdf.set_xy(x + 3, y + 38)
        pdf.cell(sticker_w - 6, 5, f"Date: ____________", 0, 1, 'L')

    buffer = io.BytesIO()
    pdf_content = pdf.output(dest='S')
    if isinstance(pdf_content, str):
        pdf_content = pdf_content.encode('latin-1')
    buffer.write(pdf_content)
    buffer.seek(0)
    return buffer
