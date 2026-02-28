import os
import re
import json
import csv
import hashlib
import base64
import logging
import threading
from typing import List, Dict, Optional
import pymupdf  # fits
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_text_splitters import RecursiveCharacterTextSplitter
from openai import OpenAI
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings

logger = logging.getLogger(__name__)

from pathlib import Path
_DATA_DIR = str(Path(__file__).resolve().parent.parent / "data")


# ─── OCR Cache ──────────────────────────────────────────────────────────────
class OCRCache:
    """
    Manages per-case OCR text caching and full-text search.

    Storage layout:
        data/cases/<case_id>/ocr_cache/
            manifest.json          ← {file_key: {status, word_count, timestamp, filename}}
            <hash>.txt             ← raw concatenated text per file (for search)
    """

    # Class-level lock for manifest I/O (shared across all instances of same path)
    _manifest_locks: Dict[str, threading.Lock] = {}
    _manifest_locks_guard = threading.Lock()

    def __init__(self, case_dir: str):
        self.cache_dir = os.path.join(case_dir, "ocr_cache")
        os.makedirs(self.cache_dir, exist_ok=True)
        self._manifest_path = os.path.join(self.cache_dir, "manifest.json")
        # Per-path lock ensures multiple OCRCache instances for the same case
        # don't corrupt each other's manifest writes
        with OCRCache._manifest_locks_guard:
            if self._manifest_path not in OCRCache._manifest_locks:
                OCRCache._manifest_locks[self._manifest_path] = threading.Lock()
        self._lock = OCRCache._manifest_locks[self._manifest_path]
        self._manifest: Dict = self._load_manifest()

    # ── Manifest I/O (thread-safe) ────────────────────────────────────
    def _load_manifest(self) -> Dict:
        with self._lock:
            if os.path.exists(self._manifest_path):
                try:
                    with open(self._manifest_path, "r", encoding="utf-8") as f:
                        return json.load(f)
                except Exception:
                    return {}
            return {}

    def _save_manifest(self):
        with self._lock:
            try:
                tmp_path = self._manifest_path + ".tmp"
                with open(tmp_path, "w", encoding="utf-8") as f:
                    json.dump(self._manifest, f, indent=2)
                os.replace(tmp_path, self._manifest_path)
            except Exception as e:
                logger.warning("Failed to save OCR manifest: %s", e)

    def reload_manifest(self):
        """Reload manifest from disk (call before reads that need fresh data)."""
        self._manifest = self._load_manifest()

    # ── Key helpers ──────────────────────────────────────────────────
    @staticmethod
    def file_key(filename: str, file_size: int) -> str:
        return f"{filename}:{file_size}"

    @staticmethod
    def _hash_key(file_key: str) -> str:
        return hashlib.md5(file_key.encode()).hexdigest()

    def _text_path(self, file_key: str) -> str:
        return os.path.join(self.cache_dir, f"{self._hash_key(file_key)}.txt")

    # ── Status API ───────────────────────────────────────────────────
    def get_status(self, file_key: str) -> Optional[str]:
        entry = self._manifest.get(file_key)
        return entry.get("status") if entry else None

    def set_status(self, file_key: str, status: str, filename: str, word_count: int = 0):
        from datetime import datetime as _dt
        entry = self._manifest.get(file_key, {})
        entry.update({
            "status": status,
            "filename": filename,
            "word_count": word_count,
            "timestamp": _dt.now().isoformat(),
        })
        self._manifest[file_key] = entry
        self._save_manifest()

    def set_skipped(self, file_key: str, filename: str):
        self.set_status(file_key, "skipped", filename, 0)

    def get_all_statuses(self) -> Dict:
        return dict(self._manifest)

    # ── Text storage ─────────────────────────────────────────────────
    def store_text(self, file_key: str, full_text: str, filename: str):
        txt_path = self._text_path(file_key)
        try:
            # Atomic write: tmp file + os.replace to prevent truncated files on crash
            tmp_path = txt_path + ".tmp"
            with open(tmp_path, "w", encoding="utf-8") as f:
                f.write(full_text)
            os.replace(tmp_path, txt_path)
        except Exception:
            pass
        word_count = len(full_text.split()) if full_text else 0
        self.set_status(file_key, "done", filename, word_count)

    def get_text(self, file_key: str) -> Optional[str]:
        txt_path = self._text_path(file_key)
        if os.path.exists(txt_path):
            try:
                with open(txt_path, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception:
                return None
        return None

    # ── Page-level storage (for resumable OCR) ─────────────────────
    def _page_text_path(self, file_key: str, page_num: int) -> str:
        return os.path.join(self.cache_dir, f"{self._hash_key(file_key)}_p{page_num}.txt")

    def set_in_progress(self, file_key: str, filename: str, pages_done: int = 0, total_pages: int = 0):
        """Mark a file as currently being OCR'd (supports resumption)."""
        from datetime import datetime as _dt
        entry = self._manifest.get(file_key, {})
        entry.update({
            "status": "in_progress",
            "filename": filename,
            "pages_done": pages_done,
            "total_pages": total_pages,
            "timestamp": _dt.now().isoformat(),
        })
        self._manifest[file_key] = entry
        self._save_manifest()

    def get_pages_done(self, file_key: str) -> int:
        """Return number of pages already OCR'd for a file (for resumption)."""
        entry = self._manifest.get(file_key, {})
        return entry.get("pages_done", 0)

    def get_total_pages(self, file_key: str) -> int:
        """Return total page count for a file."""
        entry = self._manifest.get(file_key, {})
        return entry.get("total_pages", 0)

    def store_page_text(self, file_key: str, page_num: int, text: str):
        """Store OCR text for a single page (incremental). Atomic write."""
        page_path = self._page_text_path(file_key, page_num)
        try:
            tmp_path = page_path + ".tmp"
            with open(tmp_path, "w", encoding="utf-8") as f:
                f.write(text)
            os.replace(tmp_path, page_path)
        except Exception:
            pass

    def get_page_text(self, file_key: str, page_num: int) -> Optional[str]:
        """Retrieve text for a single page."""
        page_path = self._page_text_path(file_key, page_num)
        if os.path.exists(page_path):
            try:
                with open(page_path, "r", encoding="utf-8") as f:
                    return f.read()
            except Exception:
                return None
        return None

    def finalize_file(self, file_key: str, filename: str, total_pages: int):
        """Concatenate all page texts into the final .txt file and set status to 'done'."""
        pages = []
        for p in range(total_pages):
            text = self.get_page_text(file_key, p)
            if text:
                pages.append(text)
        full_text = "\n\n".join(pages)
        if full_text.strip():
            self.store_text(file_key, full_text, filename)
        else:
            self.set_skipped(file_key, filename)
        # Clean up page files
        for p in range(total_pages):
            page_path = self._page_text_path(file_key, p)
            try:
                if os.path.exists(page_path):
                    os.remove(page_path)
            except Exception:
                pass

    # ── Full-text search ─────────────────────────────────────────────
    def search(self, query: str, max_snippets_per_file: int = 3, context_chars: int = 120) -> List[Dict]:
        """
        Search all cached text files for query (case-insensitive substring).
        Returns list of {filename, file_key, word_count, snippets: [{text, position}]}.
        """
        if not query or len(query) < 2:
            return []

        results = []
        query_lower = query.lower()

        for fkey, meta in self._manifest.items():
            if meta.get("status") != "done":
                continue
            text = self.get_text(fkey)
            if not text:
                continue

            text_lower = text.lower()
            snippets = []
            start_pos = 0
            while len(snippets) < max_snippets_per_file:
                idx = text_lower.find(query_lower, start_pos)
                if idx == -1:
                    break
                # Extract surrounding context
                snip_start = max(0, idx - context_chars)
                snip_end = min(len(text), idx + len(query) + context_chars)
                snippet_text = text[snip_start:snip_end]
                # Add ellipsis if truncated
                if snip_start > 0:
                    snippet_text = "\u2026" + snippet_text
                if snip_end < len(text):
                    snippet_text = snippet_text + "\u2026"
                snippets.append({"text": snippet_text, "position": idx})
                start_pos = idx + len(query)

            if snippets:
                results.append({
                    "filename": meta.get("filename", fkey),
                    "file_key": fkey,
                    "word_count": meta.get("word_count", 0),
                    "snippets": snippets,
                })

        # Sort by number of matches (most first)
        results.sort(key=lambda r: len(r["snippets"]), reverse=True)
        return results


class DocumentIngester:
    # Class-level per-case locks for vectorstore operations.
    # Prevents concurrent ingestion threads from colliding when
    # adding documents to the same case's Chroma collection.
    _vs_locks: Dict[str, threading.Lock] = {}
    _vs_locks_guard = threading.Lock()

    @classmethod
    def _get_vs_lock(cls, case_id: str) -> threading.Lock:
        """Get or create a per-case lock for vectorstore operations."""
        with cls._vs_locks_guard:
            if case_id not in cls._vs_locks:
                cls._vs_locks[case_id] = threading.Lock()
            return cls._vs_locks[case_id]

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            add_start_index=True,
        )
        # Legal document heading patterns
        self._heading_pattern = re.compile(
            r'^\s*(?:'
            r'(?:SECTION|ARTICLE|PART|CHAPTER|COUNT|CHARGE)\s+[IVXLCDM\d]+'
            r'|[IVXLCDM]+\.\s+[A-Z]'
            r'|\d+\.\s+[A-Z]'
            r'|[A-Z][A-Z\s]{4,}$'
            r'|#{1,4}\s+'
            r')',
            re.MULTILINE
        )

    def get_vector_store(self, case_id: str, force_local: bool = False):
        """
        Returns a Chroma vector store for a specific case.
        """
        # specialized persistence directory
        persist_dir = os.path.join(_DATA_DIR, "vector_store")

        # Check if this case previously fell back to local embeddings
        case_dir = os.path.join(_DATA_DIR, "cases", case_id)
        local_flag_path = os.path.join(case_dir, ".use_local_embeddings")
        if os.path.exists(local_flag_path):
            force_local = True

        # Use HuggingFace embeddings for local, free execution
        # Or OpenAI if available (better quality) - unless force_local is True
        if os.environ.get("OPENAI_API_KEY") and not force_local:
            try:
                from langchain_openai import OpenAIEmbeddings
                embeddings = OpenAIEmbeddings()
            except ImportError:
                embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        else:
            embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

        return Chroma(
            collection_name=f"case_{case_id}",
            embedding_function=embeddings,
            persist_directory=persist_dir
        )

    def add_to_vectorstore(self, case_id: str, documents: List[Document]):
        """
        Adds documents to the case's vector store.
        If OpenAI quota is blown, falls back to HuggingFace local models automatically.

        Thread-safe: uses a per-case lock so concurrent ingestion threads for
        different cases proceed in parallel, while same-case access is serialized.
        """
        if not documents: return

        with self._get_vs_lock(case_id):
            vectorstore = self.get_vector_store(case_id)
            try:
                vectorstore.add_documents(documents)
                logger.info(f"Added {len(documents)} documents to vector store for case {case_id}")
            except Exception as e:
                if "insufficient_quota" in str(e).lower() or "429" in str(e):
                    logger.warning(f"OpenAI Rate Limit / Quota Exceeded. Falling back to local HuggingFace Embeddings for case {case_id}...")

                    # Drop the existing empty collection to prevent dimension mismatch
                    try:
                        vectorstore.delete_collection()
                    except Exception:
                        pass

                    # Write a flag so future retrievals for this case know to use local embeddings
                    case_dir = os.path.join(_DATA_DIR, "cases", case_id)
                    os.makedirs(case_dir, exist_ok=True)
                    with open(os.path.join(case_dir, ".use_local_embeddings"), "w") as f:
                        f.write("1")

                    local_vectorstore = self.get_vector_store(case_id, force_local=True)
                    local_vectorstore.add_documents(documents)
                    logger.info(f"Recovered. Added {len(documents)} documents using local fallback embeddings.")
                else:
                    raise e

    def _detect_section_title(self, text: str) -> Optional[str]:
        """
        Detects the most prominent section heading in a block of text.
        Looks for legal document patterns: ALL CAPS headings, numbered sections,
        SECTION/ARTICLE/COUNT labels, and markdown headers.
        Returns the last detected heading or None.
        """
        lines = text.split('\n')
        last_heading = None

        for line in lines:
            stripped = line.strip()
            if not stripped or len(stripped) < 3:
                continue

            # ALL CAPS line (at least 5 chars, mostly uppercase letters)
            alpha_chars = [c for c in stripped if c.isalpha()]
            if (len(alpha_chars) >= 4 and
                sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars) > 0.8 and
                len(stripped) < 100):
                last_heading = stripped.strip(':')
                continue

            # Numbered section: "1. Something" or "I. Something"
            match = re.match(r'^\s*(?:[IVXLCDM]+|\d+)\.\s+(.+)', stripped)
            if match and len(stripped) < 100:
                last_heading = stripped
                continue

            # SECTION/ARTICLE/COUNT labels
            match = re.match(r'^\s*(?:SECTION|ARTICLE|PART|CHAPTER|COUNT|CHARGE)\s+.+', stripped, re.IGNORECASE)
            if match and len(stripped) < 100:
                last_heading = stripped
                continue

            # Markdown headers
            match = re.match(r'^#{1,4}\s+(.+)', stripped)
            if match:
                last_heading = match.group(1).strip()
                continue

        return last_heading

    def process_file(self, file_path: str, vision_model=None, force_ocr: bool = False) -> List[Document]:
        """
        Ingests a file (PDF or Image) and returns a list of LangChain Documents.

        Args:
            file_path: Path to the file.
            vision_model: A configured ChatXAI/ChatAnthropic instance for vision tasks.
            force_ocr: If True, bypass embedded text and use vision OCR on every PDF page.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()

        if file_ext == ".pdf":
            return self._process_pdf(file_path, vision_model, force_ocr=force_ocr)
        elif file_ext in [".jpg", ".jpeg", ".png"]:
            return self._process_image(file_path, vision_model)
        elif file_ext in [".mp3", ".wav", ".m4a", ".mp4", ".mpeg", ".mpga", ".webm", ".avi", ".mov", ".mkv", ".ogg", ".flac", ".aac"]:
            return self._process_media(file_path)
        elif file_ext == ".docx":
             return self._process_docx(file_path)
        elif file_ext == ".txt":
             return self._process_txt(file_path)
        elif file_ext in [".xlsx", ".xls", ".csv"]:
             return self._process_spreadsheet(file_path)

        raise ValueError(f"Unsupported file type: {file_ext}")

    def process_file_with_cache(self, file_path: str, ocr_cache: 'OCRCache',
                                vision_model=None, force_ocr: bool = False) -> List[Document]:
        """
        Process a file, using OCR cache when available.
        If the OCR cache has completed text for this file, builds Documents
        from cached text (fast, no vision model calls needed).
        Otherwise falls back to standard process_file().
        """
        fname = os.path.basename(file_path)
        try:
            fsize = os.path.getsize(file_path)
            file_key = OCRCache.file_key(fname, fsize)
        except OSError:
            file_key = fname

        if not force_ocr and ocr_cache.get_status(file_key) == "done":
            cached_text = ocr_cache.get_text(file_key)
            if cached_text and cached_text.strip():
                logger.info(f"Using cached OCR text for {fname} ({len(cached_text)} chars)")
                return self.text_splitter.create_documents(
                    [cached_text],
                    metadatas=[{
                        "source": fname,
                        "file_path": file_path,
                        "from_ocr_cache": True,
                    }]
                )

        # Cache miss — fall back to full processing
        return self.process_file(file_path, vision_model=vision_model, force_ocr=force_ocr)

    def _transcribe_image_content(self, image_data: str, media_type: str, vision_model) -> str:
        """
        Helper to send image data to LLM for transcription.
        Uses a timeout to prevent hanging on unresponsive vision models.
        """
        if not vision_model:
            return ""

        import concurrent.futures

        logger.info(f"Sending image to Vision Model ({media_type})...")
        message = HumanMessage(
            content=[
                {"type": "text", "text": "Transcribe the text in this legal document image verbatim. Maintain the layout as best as possible. If it is handwriting, transcribe it exactly. Do not summarize. Output ONLY the text."},
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{media_type};base64,{image_data}"},
                },
            ]
        )
        try:
            with concurrent.futures.ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(vision_model.invoke, [message])
                response = future.result(timeout=120)  # 2-minute timeout per page
            return response.content
        except concurrent.futures.TimeoutError:
            logger.warning("Vision model timed out after 120s -- skipping this page")
            return "[OCR timeout -- page skipped]"
        except Exception as e:
            logger.warning(f"Error calling Vision Model: {e}")
            return ""

    def _process_image(self, file_path: str, vision_model) -> List[Document]:
        """
        Uses a Multimodal LLM to transcribe text from an image.
        """
        if not vision_model:
            raise RuntimeError(
                f"No vision model configured. Cannot process image '{os.path.basename(file_path)}'. "
                "Please set an API key (xAI, Anthropic, or OpenAI) in the sidebar."
            )

        with open(file_path, "rb") as image_file:
            image_data = base64.b64encode(image_file.read()).decode("utf-8")

        # Determine media type
        ext = os.path.splitext(file_path)[1].lower().replace(".", "")
        if ext == "jpg": ext = "jpeg"
        media_type = f"image/{ext}"

        text = self._transcribe_image_content(image_data, media_type, vision_model)

        if not text:
            raise RuntimeError(
                f"Vision model returned no text for image '{os.path.basename(file_path)}'. "
                "The image may be blank or unreadable."
            )

        filename = os.path.basename(file_path)
        return self.text_splitter.create_documents(
            [text],
            metadatas=[{
                "source": filename,
                "page": 1,
                "file_path": file_path,
                "type": "image_transcription"
            }]
        )

    def _process_media(self, file_path: str) -> List[Document]:
        """
        Transcribes audio/video file using OpenAI Whisper API.
        Requires OPENAI_API_KEY env var or client init.
        Supports: mp3, wav, m4a, mp4, mpeg, mpga, webm, avi, mov, mkv, ogg, flac, aac
        """
        api_key = os.environ.get("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError(
                f"No OpenAI API key set. Cannot transcribe media file '{os.path.basename(file_path)}'. "
                "Please set OPENAI_API_KEY in the sidebar to enable audio/video transcription."
            )

        client = OpenAI(api_key=api_key)

        try:
            # Check file size (limit 25MB for API)
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            if file_size_mb > 25:
                logger.warning(f"File {file_path} is {file_size_mb:.1f}MB (Whisper API limit is 25MB). "
                      f"Consider compressing or splitting the file. Attempting anyway...")

            file_ext = os.path.splitext(file_path)[1].lower()
            media_type = "video" if file_ext in [".mp4", ".mpeg", ".webm", ".avi", ".mov", ".mkv"] else "audio"

            logger.info(f"Transcribing {media_type}: {file_path} ({file_size_mb:.1f}MB)...")

            with open(file_path, "rb") as audio_file:
                # Use verbose_json for timestamps when possible
                try:
                    transcription = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file,
                        response_format="verbose_json",
                        timestamp_granularities=["segment"]
                    )
                    # Extract text and timestamp data
                    text = transcription.text if hasattr(transcription, 'text') else str(transcription)
                    segments = []
                    if hasattr(transcription, 'segments') and transcription.segments:
                        segments = [
                            {"start": s.start, "end": s.end, "text": s.text}
                            for s in transcription.segments
                        ]
                except Exception:
                    # Fallback to plain text format
                    audio_file.seek(0)
                    transcription = client.audio.transcriptions.create(
                        model="whisper-1",
                        file=audio_file,
                        response_format="text"
                    )
                    text = str(transcription)
                    segments = []

            metadata = {
                "source": os.path.basename(file_path),
                "page": 1,
                "file_path": file_path,
                "type": "media_transcription",
                "media_type": media_type,
                "file_size_mb": round(file_size_mb, 1),
            }
            if segments:
                metadata["segments"] = json.dumps(segments)
                metadata["duration_seconds"] = segments[-1].get("end", 0) if segments else 0

            return self.text_splitter.create_documents(
                [text],
                metadatas=[metadata]
            )
        except Exception as e:
            raise RuntimeError(f"Error transcribing media '{os.path.basename(file_path)}': {e}") from e

    def _process_docx(self, file_path: str) -> List[Document]:
        """
        Extracts text from a .docx file using python-docx.
        Includes paragraph text and table cell text.
        """
        import docx
        filename = os.path.basename(file_path)

        try:
            doc = docx.Document(file_path)
            parts = []

            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        parts.append(row_text)

            text = "\n".join(parts)
            if not text.strip():
                return []

            return self.text_splitter.create_documents(
                [text],
                metadatas=[{
                    "source": filename,
                    "page": 1,
                    "file_path": file_path,
                    "type": "docx",
                    "section_title": self._detect_section_title(text),
                }]
            )
        except Exception as e:
            logger.warning(f"Error processing DOCX {file_path}: {e}")
            raise e

    def _process_txt(self, file_path: str) -> List[Document]:
        """
        Reads a plain text file.
        """
        filename = os.path.basename(file_path)
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            if not text.strip():
                return []
            return self.text_splitter.create_documents(
                [text],
                metadatas=[{
                    "source": filename,
                    "page": 1,
                    "file_path": file_path,
                    "type": "text",
                    "section_title": self._detect_section_title(text),
                }]
            )
        except Exception as e:
            logger.warning(f"Error processing TXT {file_path}: {e}")
            raise e

    def _assess_text_quality(self, text: str) -> Dict:
        """
        Scores extracted PDF text quality (0-100).
        Low scores indicate garbled/bad OCR text from scanners.
        Returns: {score, reasons}
        """
        if not text or not text.strip():
            return {"score": 0, "reasons": ["empty"]}

        reasons = []
        penalties = 0
        stripped = text.strip()
        words = stripped.split()

        if len(words) < 3:
            return {"score": 10, "reasons": ["too_few_words"]}

        # 1. Dictionary-like word ratio: words made of normal chars [a-zA-Z0-9'-.,]
        normal_word_pattern = re.compile(r"^[a-zA-Z0-9'\-.,;:!?()\"]+$")
        normal_count = sum(1 for w in words if normal_word_pattern.match(w))
        word_ratio = normal_count / max(len(words), 1)
        if word_ratio < 0.5:
            penalties += 40
            reasons.append(f"low_word_ratio:{word_ratio:.2f}")
        elif word_ratio < 0.7:
            penalties += 20
            reasons.append(f"med_word_ratio:{word_ratio:.2f}")
        elif word_ratio < 0.85:
            penalties += 8
            reasons.append(f"slight_word_ratio:{word_ratio:.2f}")

        # 2. Non-printable / unusual character density
        non_printable = sum(1 for c in stripped if ord(c) < 32 and c not in '\n\r\t')
        non_print_ratio = non_printable / max(len(stripped), 1)
        if non_print_ratio > 0.05:
            penalties += 25
            reasons.append("high_control_chars")
        elif non_print_ratio > 0.02:
            penalties += 10
            reasons.append("some_control_chars")

        # 3. Repetitive character runs (e.g., "lllll", "|||||")
        rep_runs = re.findall(r'(.)\1{3,}', stripped)
        if len(rep_runs) > 3:
            penalties += 20
            reasons.append(f"repetitive_runs:{len(rep_runs)}")
        elif len(rep_runs) > 1:
            penalties += 8
            reasons.append(f"some_rep_runs:{len(rep_runs)}")

        # 4. Digit/symbol soup — long runs of non-letter chars (not spaces)
        symbol_runs = re.findall(r'[^a-zA-Z\s]{8,}', stripped)
        if len(symbol_runs) > 2:
            penalties += 15
            reasons.append(f"symbol_soup:{len(symbol_runs)}")

        # 5. Average word length sanity (garbled text often has very long "words")
        avg_word_len = sum(len(w) for w in words) / max(len(words), 1)
        if avg_word_len > 15:
            penalties += 15
            reasons.append(f"long_avg_word:{avg_word_len:.1f}")
        elif avg_word_len < 1.5:
            penalties += 10
            reasons.append(f"short_avg_word:{avg_word_len:.1f}")

        score = max(0, min(100, 100 - penalties))
        if not reasons:
            reasons.append("clean")

        return {"score": score, "reasons": reasons}

    def _render_page_image(self, page, zoom: int = 2) -> str:
        """Renders a PDF page to a base64-encoded PNG string."""
        pix = page.get_pixmap(matrix=pymupdf.Matrix(zoom, zoom))
        return base64.b64encode(pix.tobytes("png")).decode("utf-8")

    def _process_pdf(self, file_path: str, vision_model=None, force_ocr: bool = False) -> List[Document]:
        """
        Extracts text from PDF with quality assessment.
        - force_ocr=True: skip text extraction, use vision OCR on every page.
        - Otherwise: extract text, assess quality, auto-OCR or flag for review.

        Returns List[Document]. Each doc's metadata includes:
            text_quality_score, needs_review, original_text, ocr_text

        Also populates self.ocr_reviews (list of pages flagged for human review).
        """
        docs = []
        filename = os.path.basename(file_path)
        if not hasattr(self, 'ocr_reviews'):
            self.ocr_reviews = []

        try:
            # Use context manager to guarantee file handle release even on exception
            with pymupdf.open(file_path) as doc:
                for page_num, page in enumerate(doc):
                    is_scanned = False
                    quality = {"score": 100, "reasons": ["clean"]}
                    needs_review = False
                    original_text = ""
                    ocr_text = ""

                    if force_ocr:
                        # Force OCR: skip text extraction entirely
                        original_text = page.get_text().strip()
                        if vision_model:
                            logger.info(f"Force OCR: Page {page_num+1} -- sending to Vision Model...")
                            image_data = self._render_page_image(page, zoom=2)
                            ocr_text = self._transcribe_image_content(image_data, "image/png", vision_model)
                            text = ocr_text if ocr_text else original_text
                            is_scanned = True
                            quality = self._assess_text_quality(text)
                        else:
                            logger.warning(f"Force OCR enabled but no Vision Model available. Using embedded text for page {page_num+1}.")
                            text = original_text
                    else:
                        # Normal mode: extract text, assess quality
                        text = page.get_text()
                        original_text = text.strip()

                        if len(text.strip()) < 50:
                            # Very sparse text — likely a scanned page
                            quality = {"score": 5, "reasons": ["sparse_text"]}
                            if vision_model:
                                logger.info(f"Page {page_num+1} appears scanned (sparse text). OCR'ing with Vision Model...")
                                image_data = self._render_page_image(page, zoom=2)
                                ocr_text = self._transcribe_image_content(image_data, "image/png", vision_model)
                                if ocr_text:
                                    text = ocr_text
                                    is_scanned = True
                                    quality = self._assess_text_quality(text)
                            else:
                                logger.warning(f"Page {page_num+1} is scanned but no Vision Model available.")
                        else:
                            # Has substantial text — check quality
                            quality = self._assess_text_quality(text)

                            if quality["score"] < 60:
                                # Poor quality — auto-OCR if possible
                                if vision_model:
                                    logger.info(f"Page {page_num+1} has low-quality text (score={quality['score']}). Auto-OCR'ing...")
                                    image_data = self._render_page_image(page, zoom=2)
                                    ocr_text = self._transcribe_image_content(image_data, "image/png", vision_model)
                                    if ocr_text:
                                        text = ocr_text
                                        is_scanned = True
                                        quality = self._assess_text_quality(text)
                                else:
                                    logger.warning(f"Page {page_num+1} has low-quality text (score={quality['score']}) but no Vision Model.")
                                    needs_review = True
                            elif quality["score"] < 75:
                                # Medium quality — flag for human review
                                needs_review = True
                                if vision_model:
                                    logger.info(f"Page {page_num+1} has medium-quality text (score={quality['score']}). Getting OCR for comparison...")
                                    image_data = self._render_page_image(page, zoom=2)
                                    ocr_text = self._transcribe_image_content(image_data, "image/png", vision_model)

                    if not text.strip():
                        continue

                    # Flag for review queue if needed
                    if needs_review:
                        review_entry = {
                            "id": hashlib.md5(f"{filename}_{page_num}".encode()).hexdigest()[:8],
                            "source": filename,
                            "file_path": file_path,
                            "page": page_num + 1,
                            "total_pages": doc.page_count,
                            "extracted_text": original_text[:2000],
                            "ocr_text": (ocr_text or "")[:2000],
                            "quality_score": quality["score"],
                            "quality_reasons": quality["reasons"],
                            "status": "pending",
                            "resolution": None,
                            "replacement_text": None,
                        }
                        self.ocr_reviews.append(review_entry)

                    page_docs = self.text_splitter.create_documents(
                        [text],
                        metadatas=[{
                            "source": filename,
                            "page": page_num + 1,
                            "total_pages": doc.page_count,
                            "file_path": file_path,
                            "ocr_applied": is_scanned,
                            "text_quality_score": quality["score"],
                            "needs_review": needs_review,
                            "section_title": self._detect_section_title(text),
                        }]
                    )
                    docs.extend(page_docs)
        except Exception as e:
            logger.warning(f"Error processing PDF {file_path}: {e}")
            raise e

        return docs

    def _process_spreadsheet(self, file_path: str) -> List[Document]:
        """
        Extracts data from Excel (.xlsx/.xls) or CSV files.
        Each sheet becomes a separate Document with pipe-delimited table format.
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        sheets_data = []  # list of (sheet_name, headers, rows)

        try:
            if file_ext == ".csv":
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    reader = csv.reader(f)
                    all_rows = list(reader)
                if all_rows:
                    headers = all_rows[0]
                    data_rows = all_rows[1:]
                    sheets_data.append(("Sheet1", headers, data_rows))
            else:
                # Excel files
                try:
                    from openpyxl import load_workbook
                    wb = load_workbook(file_path, read_only=True, data_only=True)
                    for sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        rows = []
                        headers = []
                        for i, row in enumerate(ws.iter_rows(values_only=True)):
                            str_row = [str(cell) if cell is not None else "" for cell in row]
                            if i == 0:
                                headers = str_row
                            else:
                                rows.append(str_row)
                        if headers or rows:
                            sheets_data.append((sheet_name, headers, rows))
                    wb.close()
                except ImportError:
                    logger.warning("openpyxl not installed. Install with: pip install openpyxl")
                    return [
                        Document(
                            page_content=f"[Excel file: {os.path.basename(file_path)} -- openpyxl required for processing]",
                            metadata={"source": os.path.basename(file_path), "page": 1, "type": "spreadsheet", "error": "openpyxl_missing"}
                        )
                    ]
        except Exception as e:
            raise RuntimeError(f"Error processing spreadsheet '{os.path.basename(file_path)}': {e}") from e

        documents = []
        for sheet_name, headers, rows in sheets_data:
            # Build pipe-delimited table text
            lines = []
            if headers:
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows:
                # Pad/trim row to match header count
                padded = row + [""] * max(0, len(headers) - len(row))
                padded = padded[:max(len(headers), len(row))]
                lines.append("| " + " | ".join(padded) + " |")

            table_text = "\n".join(lines)

            metadata = {
                "source": os.path.basename(file_path),
                "page": 1,
                "file_path": file_path,
                "type": "spreadsheet",
                "sheet_name": sheet_name,
                "row_count": len(rows),
                "column_count": len(headers),
                "headers": json.dumps(headers),
            }

            # For large spreadsheets, split into chunks
            if len(table_text) > 2000:
                split_docs = self.text_splitter.create_documents(
                    [table_text],
                    metadatas=[metadata]
                )
                documents.extend(split_docs)
            else:
                documents.append(
                    Document(page_content=table_text, metadata=metadata)
                )

        if not documents:
            documents.append(
                Document(
                    page_content=f"[Empty spreadsheet: {os.path.basename(file_path)}]",
                    metadata={"source": os.path.basename(file_path), "page": 1, "type": "spreadsheet", "empty": True}
                )
            )

        return documents

    def calculate_tokens(self, docs: List[Document]) -> int:
        """
        Rough estimation of tokens for user warning.
        """
        total_chars = sum(len(d.page_content) for d in docs)
        return total_chars // 4  # Rough approximation

# ---- Auto-Classification (zero LLM cost) ----------------------------------
# Pattern-matched document tagging based on filename and first-page content.

import re as _re

# Filename pattern → tag mapping (case-insensitive)
_FILENAME_TAG_PATTERNS = {
    r"police|incident\s*report|offense\s*report|arrest|affidavit": "Police Report",
    r"financial|bank|account|invoice|receipt|ledger|tax": "Financial Records",
    r"witness|witness\s*statement|sworn\s*statement": "Witness Statement",
    r"medical|hospital|doctor|clinic|injury|diagnosis|treatment": "Medical Records",
    r"photo|video|image|bodycam|dashcam|surveillance|screenshot": "Photos/Video",
    r"court|filing|order|motion|complaint|indictment|subpoena|plea|judgment|docket": "Court Filing",
    r"expert|forensic|lab\s*report|toxicology|dna|ballistic": "Expert Report",
    r"letter|email|correspondence|memo|fax": "Correspondence",
    r"contract|agreement|lease|deed|settlement|mou": "Contract/Agreement",
    r"deposition|depo\b": "Deposition",
    r"discovery|interrogator|request\s*for\s*production|rfa|rfp|rog": "Discovery",
}

# First-page content header patterns → tag mapping
_CONTENT_TAG_PATTERNS = {
    r"POLICE\s*DEPARTMENT|LAW\s*ENFORCEMENT|INCIDENT\s*REPORT|OFFENSE\s*REPORT|ARREST\s*REPORT": "Police Report",
    r"WITNESS\s*STATEMENT|SWORN\s*STATEMENT|AFFIDAVIT\s*OF": "Witness Statement",
    r"MEDICAL\s*RECORD|PATIENT\s*NAME|DIAGNOSIS|DISCHARGE\s*SUMMARY|HOSPITAL": "Medical Records",
    r"FINANCIAL\s*STATEMENT|BANK\s*STATEMENT|ACCOUNT\s*SUMMARY": "Financial Records",
    r"COURT\s*OF|IN\s*THE\s*MATTER\s*OF|SUPERIOR\s*COURT|CIRCUIT\s*COURT|DOCKET\s*NO|CASE\s*NO": "Court Filing",
    r"EXPERT\s*REPORT|FORENSIC\s*ANALYSIS|LABORATORY\s*REPORT|TOXICOLOGY": "Expert Report",
    r"DEPOSITION\s*OF|EXAMINATION\s*BEFORE\s*TRIAL": "Deposition",
    r"INTERROGATOR|REQUEST\s*FOR\s*PRODUCTION|REQUEST\s*FOR\s*ADMISSION": "Discovery",
}


def auto_classify_file(filename: str, first_page_text: str = "") -> Optional[str]:
    """
    Classify a file by filename patterns and first-page content keywords.
    Returns a single tag string, or None if no match.
    Zero LLM cost — pure regex matching.
    """
    if not filename:
        return None

    name_lower = filename.lower()

    # 1. Check filename patterns first (more reliable)
    for pattern, tag in _FILENAME_TAG_PATTERNS.items():
        if _re.search(pattern, name_lower):
            return tag

    # 2. Check file extension for images/video
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"):
        return "Photos/Video"
    if ext in (".mp4", ".avi", ".mov", ".mkv", ".webm"):
        return "Photos/Video"

    # 3. Check first-page content headers
    if first_page_text:
        header = first_page_text[:500]
        for pattern, tag in _CONTENT_TAG_PATTERNS.items():
            if _re.search(pattern, header, _re.IGNORECASE):
                return tag

    return None


def smart_split_pdf(file_path: str, min_pages_per_section: int = 3) -> list:
    """
    Analyze a PDF and detect logical section boundaries.

    Returns list of sections:
        [{"pages": [0,1,2], "title": "Police Report", "suggested_tag": "Police Report"}, ...]

    Section boundaries detected by:
    - Large font headers (>14pt)
    - Page content starting with EXHIBIT / ATTACHMENT labels
    - Significant whitespace gaps with new headers
    """
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF not installed; cannot split PDF")
        return []

    try:
        doc = fitz.open(file_path)
    except Exception as e:
        logger.warning("Could not open PDF for splitting: %s", e)
        return []

    total_pages = len(doc)
    if total_pages < 5:
        doc.close()
        return []  # Too small to split

    # Detect section boundaries
    boundaries = []  # (page_idx, title)
    _section_patterns = _re.compile(
        r"(?:EXHIBIT\s+[A-Z0-9]+|ATTACHMENT\s+[A-Z0-9]+|APPENDIX\s+[A-Z0-9]+|"
        r"TAB\s+\d+|SECTION\s+\d+|PART\s+[IVX\d]+)",
        _re.IGNORECASE,
    )

    for page_num in range(total_pages):
        try:
            page = doc[page_num]
            text = page.get_text("text", sort=True)
            first_lines = text.strip()[:200] if text else ""

            # Check for section markers in first few lines
            match = _section_patterns.search(first_lines)
            if match and page_num > 0:
                boundaries.append((page_num, match.group().strip()))
                continue

            # Check for large font headers (>14pt)
            blocks = page.get_text("dict", sort=True).get("blocks", [])
            for block in blocks[:3]:  # Only check first 3 blocks on page
                for line in block.get("lines", [])[:2]:
                    for span in line.get("spans", []):
                        if span.get("size", 0) > 14 and span.get("text", "").strip():
                            header_text = span["text"].strip()
                            if len(header_text) > 3 and page_num > 0:
                                boundaries.append((page_num, header_text[:60]))
                                break
        except Exception:
            continue

    doc.close()

    if not boundaries:
        return []  # No clear section boundaries found

    # Build sections from boundaries
    sections = []
    boundary_pages = [b[0] for b in boundaries]
    boundary_titles = [b[1] for b in boundaries]

    # First section: page 0 to first boundary
    if boundary_pages[0] > 0:
        sections.append({
            "pages": list(range(0, boundary_pages[0])),
            "title": "Opening Section",
            "suggested_tag": auto_classify_file("", "") or "",
        })

    for i, (start_page, title) in enumerate(boundaries):
        end_page = boundaries[i + 1][0] if i + 1 < len(boundaries) else total_pages
        pages = list(range(start_page, end_page))
        if len(pages) < min_pages_per_section and sections:
            # Merge small sections into previous
            sections[-1]["pages"].extend(pages)
        else:
            suggested_tag = auto_classify_file(title) or ""
            sections.append({
                "pages": pages,
                "title": title,
                "suggested_tag": suggested_tag,
            })

    return sections


if __name__ == "__main__":
    # Test stub
    ingester = DocumentIngester()
    logger.info("DocumentIngester initialized.")
