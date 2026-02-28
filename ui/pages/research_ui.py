"""Research & Draft UI module -- ported from ui_modules/research_ui.py."""
import logging
import os
import io
import json

import pandas as pd
import streamlit as st
from datetime import datetime, date

from core.nodes import (
    generate_investigation_plan, generate_draft_document,
    analyze_media_forensic, analyze_spreadsheet, generate_cheat_sheet,
    generate_lexis_queries, analyze_lexis_results,
    generate_demand_letter, conduct_legal_research,
)
from core.append_only import safe_update_and_save
from core.cost_tracker import format_cost_badge
from core.export import format_court_document, get_jurisdiction_list

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Main render
# ---------------------------------------------------------------------------

def render(case_id, case_mgr, results, tabs, selected_group, nav_groups, model_provider, prep_id):
    """Render the Research & Draft UI based on active_tab."""
    estimate_cost = lambda text, mp: format_cost_badge(text, mp)

    with tabs[0]:  # INVESTIGATION
        st.subheader("Investigation Action Plan")

        # Manual Task Entry
        with st.expander("Add New Investigation Task", expanded=False):
            with st.form("new_task_form"):
                c1, c2, c3 = st.columns([3, 2, 1])
                with c1: new_action = st.text_input("Action Item")
                with c2: new_reason = st.text_input("Reasoning/Source")
                with c3: new_prio = st.selectbox("Priority", ["High", "Medium", "Low"])

                if st.form_submit_button("Add Task"):
                    if new_action:
                        new_task = {
                            "action": new_action,
                            "reason": new_reason,
                            "priority": new_prio,
                            "status": "pending",
                            "created_at": date.today().isoformat()
                        }
                        inv_plan = st.session_state.agent_results.get("investigation_plan", [])
                        if not isinstance(inv_plan, list):
                            inv_plan = []
                        inv_plan.append(new_task)
                        st.session_state.agent_results["investigation_plan"] = inv_plan
                        case_mgr.save_prep_state(st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results)
                        st.success(f"Added: {new_action}")
                        st.rerun()

        if st.button(f"Regenerate Investigation Plan {estimate_cost(results.get('case_summary', ''), model_provider)}", key="regen_investigation"):
            with st.spinner("Regenerating Investigation Plan"):
                updates = generate_investigation_plan(st.session_state.agent_results)
                safe_update_and_save(case_mgr, st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results, updates)
                st.rerun()

        investigation = results.get("investigation_plan", [])

        i_data = []
        if isinstance(investigation, str):
            try:
                i_data = json.loads(investigation.replace("```json", "").replace("```", ""))
            except Exception:
                st.markdown(investigation)
        elif isinstance(investigation, list):
            i_data = investigation

        if i_data:
            # Sort: pending first, then completed
            _sorted_inv = sorted(i_data, key=lambda x: (0 if x.get("status", "pending") == "pending" else 1) if isinstance(x, dict) else 0)

            _total_tasks = len(_sorted_inv)
            _done_tasks = sum(1 for t in _sorted_inv if isinstance(t, dict) and t.get('status') == 'completed')
            _pct = _done_tasks / _total_tasks if _total_tasks > 0 else 0
            st.progress(_pct, text=f"{_done_tasks}/{_total_tasks} tasks completed ({_pct:.0%})")

            for idx, task in enumerate(_sorted_inv):
                if isinstance(task, dict):
                    action = task.get("action", "Unknown")
                    reason = task.get("reason", "")
                    priority = task.get("priority", "Low")
                    status = task.get("status", "pending")

                    if priority == "High":
                        prio_badge = "HIGH"
                    elif priority == "Medium":
                        prio_badge = "MED"
                    else:
                        prio_badge = "LOW"

                    is_done = status == "completed"
                    display_action = f"~~{action}~~" if is_done else action

                    tc1, tc2, tc3, tc4 = st.columns([0.05, 3, 1, 0.5])
                    with tc1:
                        _new_done = st.checkbox("done", value=is_done, key=f"_inv_chk_{idx}", label_visibility="collapsed")
                        if _new_done != is_done:
                            _sorted_inv[idx]["status"] = "completed" if _new_done else "pending"
                            st.session_state.agent_results["investigation_plan"] = _sorted_inv
                            case_mgr.save_prep_state(st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results)
                            st.rerun()
                    with tc2:
                        st.markdown(f"[{prio_badge}] {display_action}")
                        if reason:
                            st.caption(reason)
                    with tc3:
                        if task.get("created_at"):
                            st.caption(f"Added: {task.get('created_at')}")
                    with tc4:
                        if st.button("X", key=f"_inv_del_{idx}"):
                            _sorted_inv.pop(idx)
                            st.session_state.agent_results["investigation_plan"] = _sorted_inv
                            case_mgr.save_prep_state(st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results)
                            st.rerun()
        else:
            st.info("No investigation plan generated.")

    with tabs[1]:  # DRAFT DOCUMENT
        st.subheader("AI Document Drafter")
        st.caption("Generate court-ready motions, memoranda, and legal documents")

        dc1, dc2 = st.columns(2)
        with dc1:
            doc_type = st.selectbox("Document Type", [
                "Motion to Suppress", "Motion to Dismiss", "Motion in Limine",
                "Memorandum of Law", "Trial Brief", "Sentencing Memorandum",
                "Pretrial Memorandum", "Reply Brief", "Position Statement",
                "Demand Letter", "Complaint", "Answer", "Custom"
            ], key="draft_doc_type")
        with dc2:
            jurisdictions = get_jurisdiction_list()
            jurisdiction = st.selectbox("Jurisdiction", jurisdictions, key="draft_jurisdiction")

        if doc_type == "Custom":
            custom_type = st.text_input("Custom Document Type", placeholder="e.g. Motion for Continuance")

        custom_instructions = st.text_area(
            "Custom Instructions (optional)",
            placeholder="e.g. Focus on Fourth Amendment issues, emphasize the warrantless search, cite State v. Smith...",
            height=100,
            key="draft_instructions"
        )

        if st.button(
            f"Generate Draft {estimate_cost(results.get('case_summary', '') + results.get('strategy_notes', ''), model_provider)}",
            type="primary",
            key="gen_draft_doc"
        ):
            if results.get("case_summary"):
                _doc_type_final = custom_type if doc_type == "Custom" else doc_type
                with st.spinner(f"Drafting {_doc_type_final}..."):
                    _specific_ctx = ""
                    if jurisdiction:
                        _specific_ctx += f"Jurisdiction: {jurisdiction}\n"
                    if custom_instructions:
                        _specific_ctx += f"Instructions: {custom_instructions}\n"
                    draft_result = generate_draft_document(
                        st.session_state.agent_results,
                        doc_type=_doc_type_final,
                        specific_context=_specific_ctx,
                    )
                    st.session_state.agent_results["draft_document"] = draft_result.get("drafted_documents", draft_result.get("draft_document", ""))
                    case_mgr.save_prep_state(st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results)
                    st.rerun()
            else:
                st.warning("Run a case analysis first before generating legal documents.")

        draft = results.get("draft_document", "")
        if draft:
            st.divider()

            # Court formatting
            with st.expander("Court Formatting", expanded=False):
                fmt_juris = st.selectbox("Format for Court", jurisdictions, key="fmt_jurisdiction")
                if st.button("Apply Court Formatting", key="apply_fmt"):
                    formatted = format_court_document(draft, fmt_juris)
                    st.session_state.agent_results["draft_document_formatted"] = formatted
                    st.rerun()

            formatted_draft = results.get("draft_document_formatted", draft)
            st.markdown(formatted_draft)

            st.divider()
            dl1, dl2, dl3 = st.columns(3)
            with dl1:
                st.download_button(
                    "Download (.md)",
                    data=formatted_draft,
                    file_name=f"draft_{doc_type.replace(' ', '_').lower()}_{case_id}.md",
                    mime="text/markdown",
                    key="dl_draft_md"
                )
            with dl2:
                try:
                    from core.export import generate_pdf_report
                    _pdf_state = {"case_summary": formatted_draft, "strategy_notes": ""}
                    _pdf = generate_pdf_report(_pdf_state, f"{doc_type} -- {case_id}")
                    st.download_button(
                        "Download (.pdf)",
                        data=_pdf,
                        file_name=f"draft_{doc_type.replace(' ', '_').lower()}_{case_id}.pdf",
                        mime="application/pdf",
                        key="dl_draft_pdf"
                    )
                except Exception:
                    pass
            with dl3:
                try:
                    from docx import Document as _DocxDoc
                    _docx_doc = _DocxDoc()
                    _docx_doc.add_heading(doc_type, 0)
                    for _para in formatted_draft.split("\n"):
                        if _para.strip():
                            _docx_doc.add_paragraph(_para)
                    _docx_buf = io.BytesIO()
                    _docx_doc.save(_docx_buf)
                    _docx_buf.seek(0)
                    st.download_button(
                        "Download (.docx)",
                        data=_docx_buf,
                        file_name=f"draft_{doc_type.replace(' ', '_').lower()}_{case_id}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_draft_docx"
                    )
                except Exception:
                    pass
        else:
            st.info("No draft generated yet. Select a document type and click Generate above.")

    with tabs[2]:  # SPREADSHEET
        st.subheader("Spreadsheet Analysis")
        st.caption("Analyze Excel and CSV files uploaded to the case for financial data, timelines, and cross-references.")

        _spreadsheets = st.session_state.agent_results.get("spreadsheet_data", {})
        if _spreadsheets and isinstance(_spreadsheets, dict):
            for src_name, src_data in _spreadsheets.items():
                if not isinstance(src_data, dict):
                    continue
                sheet_name = src_data.get("sheet_name", "Sheet1")
                row_count = src_data.get("row_count", "?")
                col_count = src_data.get("col_count", "?")
                headers = src_data.get("headers", [])
                full_text = src_data.get("full_text", "")

                with st.expander(f"**{src_name}** -- Sheet: {sheet_name} | Rows: {row_count} | Cols: {col_count}", expanded=True):
                    if headers:
                        st.caption(f"Columns: {', '.join(headers[:15])}")

                    try:
                        lines = full_text.strip().split("\n")
                        table_lines = [l for l in lines if l.startswith("|") and "---" not in l]
                        if table_lines and headers:
                            data_rows = []
                            for line in table_lines[1:21]:
                                cells = [c.strip() for c in line.strip("|").split("|")]
                                data_rows.append(cells)
                            if data_rows:
                                max_cols = max(len(headers), max(len(r) for r in data_rows))
                                padded_headers = headers + [f"Col{i+1}" for i in range(len(headers), max_cols)]
                                padded_rows = [r + [""] * (max_cols - len(r)) for r in data_rows]
                                df = pd.DataFrame(padded_rows, columns=padded_headers[:max_cols])
                                st.dataframe(df, use_container_width=True)
                                if int(row_count) > 20:
                                    st.caption(f"Showing first 20 of {row_count} rows")
                    except Exception:
                        st.text_area("Raw Data", full_text[:3000], height=200, key=f"ss_raw_{src_name}")

                    col_a, col_b = st.columns(2)
                    with col_a:
                        if st.button(f"Analyze Spreadsheet {estimate_cost(full_text, model_provider)}", key=f"ss_analyze_{src_name}"):
                            with st.spinner("Analyzing spreadsheet data..."):
                                state = dict(st.session_state.agent_results)
                                sheet_info = f"Sheet: {sheet_name} | Rows: {row_count} | Columns: {col_count}\nHeaders: {', '.join(headers)}\n"
                                result = analyze_spreadsheet(
                                    state,
                                    full_text,
                                    source_name=src_name,
                                    sheet_info=sheet_info
                                )
                                if "spreadsheet_analysis" not in st.session_state.agent_results:
                                    st.session_state.agent_results["spreadsheet_analysis"] = {}
                                st.session_state.agent_results["spreadsheet_analysis"].update(result.get("spreadsheet_analysis", {}))
                                st.rerun()

                    with col_b:
                        st.download_button(
                            "Download Data",
                            full_text,
                            file_name=f"{src_name}_data.txt",
                            key=f"dl_ss_{src_name}"
                        )

                    ss_analysis = st.session_state.agent_results.get("spreadsheet_analysis", {}).get(src_name, {})
                    if ss_analysis and not ss_analysis.get("_parse_error"):
                        st.markdown("---")
                        st.markdown("#### Spreadsheet Analysis Results")

                        if ss_analysis.get("overall_assessment"):
                            st.info(ss_analysis["overall_assessment"])

                        stabs = st.tabs(["Summary", "Financial", "Timeline", "Key Findings", "Cross-Reference"])

                        with stabs[0]:
                            ds = ss_analysis.get("data_summary", {})
                            if ds:
                                st.markdown(f"**Description**: {ds.get('description', '')}")
                                mc1, mc2, mc3 = st.columns(3)
                                with mc1:
                                    st.metric("Records", ds.get('record_count', '?'))
                                with mc2:
                                    st.metric("Date Range", ds.get('date_range', 'N/A'))
                                with mc3:
                                    st.metric("Completeness", ds.get('completeness', 'N/A')[:20])
                                if ds.get('key_columns'):
                                    st.markdown(f"**Key Columns**: {', '.join(ds['key_columns'])}")

                        with stabs[1]:
                            fa = ss_analysis.get("financial_analysis", {})
                            if fa and fa.get("has_financial_data"):
                                if fa.get('total_amounts'):
                                    st.metric("Total Amounts", fa['total_amounts'])
                                anomalies = fa.get('anomalies', [])
                                if anomalies:
                                    st.markdown("**Anomalies:**")
                                    for an in anomalies:
                                        st.warning(f"{an.get('description', '')} -- Rows: {an.get('rows_affected', '')}")
                                        st.caption(an.get('significance', ''))
                                patterns = fa.get('patterns', [])
                                if patterns:
                                    st.markdown("**Patterns:**")
                                    for p in patterns:
                                        st.markdown(f"- {p}")
                            else:
                                st.caption("No financial data detected in this spreadsheet.")

                        with stabs[2]:
                            td = ss_analysis.get("timeline_data", {})
                            if td and td.get("has_dates"):
                                events = td.get('chronological_events', [])
                                if events:
                                    for evt in events:
                                        st.markdown(f"**{evt.get('date', '')}** -- {evt.get('event', '')}")
                                        st.caption(evt.get('significance', ''))
                                gaps = td.get('gaps', [])
                                if gaps:
                                    st.warning(f"Gaps: {', '.join(gaps)}")
                                if td.get('clustering'):
                                    st.info(f"Clustering: {td['clustering']}")
                            else:
                                st.caption("No date-based data detected.")

                        with stabs[3]:
                            findings = ss_analysis.get("key_findings", [])
                            if findings:
                                for kf in findings:
                                    priority = kf.get('priority', 'medium')
                                    st.markdown(f"[{priority.upper()}] **{kf.get('finding', '')}**")
                                    st.caption(f"Data: {kf.get('supporting_data', '')}")
                                    st.caption(f"Legal Relevance: {kf.get('legal_relevance', '')}")
                                    st.markdown("---")
                            else:
                                st.caption("No key findings identified.")

                        with stabs[4]:
                            cr = ss_analysis.get("cross_reference", {})
                            if cr:
                                matches = cr.get('matches_case_facts', [])
                                if matches:
                                    st.success("**Confirms Case Facts:**")
                                    for m in matches:
                                        st.markdown(f"- {m}")
                                contras = cr.get('contradicts_case_facts', [])
                                if contras:
                                    st.error("**Contradicts Case Facts:**")
                                    for c in contras:
                                        st.markdown(f"- {c}")
                                leads = cr.get('new_leads', [])
                                if leads:
                                    st.info("**New Leads:**")
                                    for l in leads:
                                        st.markdown(f"- {l}")
                                followup = cr.get('suggested_follow_up', [])
                                if followup:
                                    st.markdown("**Suggested Follow-Up:**")
                                    for f in followup:
                                        st.markdown(f"- {f}")
                            else:
                                st.caption("No cross-reference data.")

                    elif ss_analysis and ss_analysis.get("_parse_error"):
                        st.markdown("---")
                        st.markdown("#### Analysis (Raw)")
                        st.write(ss_analysis.get("overall_assessment", ""))
        else:
            st.info("No spreadsheets found. Upload Excel (.xlsx) or CSV files in the sidebar.")

    with tabs[3]:  # LEGAL RESEARCH
        st.subheader("Legal Research Agent")
        st.caption("Autonomous Web Research for Statutes & Case Law")

        if st.button(f"Conduct Deep Research {estimate_cost(str(results.get('case_summary', '')) + results.get('strategy_notes', ''), model_provider)}", type="primary"):
            with st.spinner("Agent is researching live on the web..."):
                updates = conduct_legal_research(st.session_state.agent_results)
                safe_update_and_save(case_mgr, st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results, updates)
                st.rerun()

        search_data = results.get("legal_research_data", [])
        summary = results.get("research_summary", "")

        if summary:
            st.markdown("### Research Memo")
            st.markdown(summary)
            st.divider()

        if search_data:
            st.markdown("### Search Sources")
            if isinstance(search_data, str):
                try:
                    search_data = json.loads(search_data)
                except (json.JSONDecodeError, ValueError):
                    logger.info("Failed to parse legal_research_data JSON")

            if isinstance(search_data, list):
                for item in search_data:
                    with st.expander(f"Query: {item.get('query', 'Unknown')}", expanded=False):
                        st.write(item.get("result"))

        if not summary and not search_data:
            st.info("No research conducted yet. Click the button above to start.")

    # LEXIS+ SEARCH ASSISTANT
    _lex_tab_items = nav_groups.get("Research & Draft", nav_groups.get("\ud83d\udcda Research & Draft", []))
    _lex_tab_idx = -2
    for _i, _label in enumerate(_lex_tab_items):
        if "Lexis" in _label:
            _lex_tab_idx = _i
            break

    with tabs[_lex_tab_idx]:  # LEXIS+ ASSISTANT
        st.subheader("Lexis+ Search Assistant")
        st.caption("Generate optimized search queries for your Lexis+ subscription, then paste results for AI analysis.")

        # Section 1: Query Generator
        st.markdown("### Query Generator")
        lex_focus = st.text_input(
            "Research focus (optional)",
            placeholder="e.g., Motion to suppress -- Fourth Amendment, chain of custody issues",
            key="_lexis_focus",
        )
        if st.button("Generate Lexis+ Queries", type="primary", key="_gen_lexis_q"):
            if results.get("case_summary"):
                with st.spinner("Generating optimized Lexis+ Boolean queries..."):
                    q_result = generate_lexis_queries(st.session_state.agent_results, research_focus=lex_focus)
                    st.session_state["lexis_queries"] = q_result.get("lexis_queries", [])
                    st.rerun()
            else:
                st.warning("Analyze case documents first to generate targeted queries.")

        lex_queries = st.session_state.get("lexis_queries", [])
        if lex_queries:
            for qi, q in enumerate(lex_queries):
                search_str = q.get("search_string", "") if isinstance(q, dict) else str(q)
                desc = q.get("description", "") if isinstance(q, dict) else ""
                filters = q.get("filters", {}) if isinstance(q, dict) else {}
                relevance = q.get("case_relevance", "") if isinstance(q, dict) else ""

                with st.container():
                    st.markdown(
                        f'<div style="background:rgba(99,102,241,0.08);border:1px solid rgba(99,102,241,0.2);'
                        f'border-radius:10px;padding:14px 18px;margin-bottom:10px;">'
                        f'<div style="font-weight:700;font-size:13px;color:#818cf8;margin-bottom:4px;">'
                        f'Query {qi+1}</div>'
                        f'<code style="font-size:13px;display:block;margin-bottom:8px;word-break:break-all;">'
                        f'{search_str}</code>'
                        f'<div style="font-size:12px;opacity:0.8;margin-bottom:4px;">{desc}</div>'
                        + (f'<div style="font-size:11px;opacity:0.6;">'
                           f'{filters.get("jurisdiction","")} | '
                           f'{filters.get("date_range","")} | '
                           f'{filters.get("court_level","")}</div>' if filters else '')
                        + (f'<div style="font-size:11px;color:#22c55e;margin-top:4px;">{relevance}</div>' if relevance else '')
                        + '</div>',
                        unsafe_allow_html=True,
                    )
                    st.code(search_str, language=None)
            st.info("Copy any query above and paste it directly into the Lexis+ search bar.")

        st.divider()

        # Section 2: Results Intake
        st.markdown("### Paste Lexis+ Results")
        st.caption("Copy case text from Lexis+ and paste it here for AI analysis.")

        lex_context_options = ["(None -- general analysis)"]
        if lex_queries:
            for qi, q in enumerate(lex_queries):
                label = q.get("search_string", str(q))[:80] if isinstance(q, dict) else str(q)[:80]
                lex_context_options.append(f"Query {qi+1}: {label}")

        lex_ctx_sel = st.selectbox("Tag which query these results are for", lex_context_options, key="_lex_ctx")

        lex_paste = st.text_area(
            "Paste Lexis+ text",
            height=250,
            placeholder="Paste the full case text, headnotes, or search results from Lexis+ here...",
            key="_lex_paste",
        )

        lex_upload = st.file_uploader(
            "Or upload a file",
            type=["txt", "docx", "pdf"],
            key="_lex_upload",
            help="Upload case text exported from Lexis+",
        )

        lex_text = lex_paste
        if lex_upload and not lex_paste:
            try:
                raw = lex_upload.read()
                lex_text = raw.decode("utf-8", errors="replace")
            except Exception:
                lex_text = ""

        query_ctx = ""
        if lex_ctx_sel and lex_ctx_sel != "(None -- general analysis)" and lex_queries:
            try:
                idx = int(lex_ctx_sel.split(":")[0].replace("Query ", "")) - 1
                q = lex_queries[idx]
                query_ctx = q.get("search_string", "") if isinstance(q, dict) else str(q)
            except Exception:
                pass

        if st.button("Analyze Results", type="primary", key="_analyze_lex", disabled=not lex_text):
            with st.spinner("Analyzing Lexis+ results -- extracting cases, holdings, and strategy..."):
                analysis = analyze_lexis_results(st.session_state.agent_results, lex_text, query_context=query_ctx)
                existing = st.session_state.get("lexis_research_results", [])
                new_entry = {
                    "query_context": query_ctx,
                    "analysis": analysis.get("lexis_analysis", {}),
                    "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M"),
                }
                existing.append(new_entry)
                st.session_state["lexis_research_results"] = existing
                st.rerun()

        st.divider()

        # Section 3: Analysis Dashboard
        st.markdown("### Analysis Dashboard")
        lex_results = st.session_state.get("lexis_research_results", [])
        if not lex_results:
            st.info("No Lexis+ results analyzed yet. Generate queries above, run them in Lexis+, then paste results here.")
        else:
            for ri, entry in enumerate(reversed(lex_results)):
                analysis = entry.get("analysis", {})
                cases = analysis.get("cases", [])
                summary = analysis.get("summary", "")
                ts = entry.get("timestamp", "")
                qc = entry.get("query_context", "General")

                with st.expander(f"Analysis {len(lex_results) - ri} -- {ts} ({len(cases)} cases found)", expanded=(ri == 0)):
                    if qc:
                        st.caption(f"Query: `{qc[:100]}`")
                    if summary:
                        st.markdown("**Research Summary**")
                        st.markdown(summary)
                        st.divider()

                    if cases:
                        table_data = []
                        for c in cases:
                            fav = c.get("favorability", "?")
                            icon = "YES" if fav == "FAVORABLE" else "NO" if fav == "UNFAVORABLE" else "NEUTRAL"
                            table_data.append({
                                "Case": c.get("citation", "Unknown"),
                                "Court": c.get("court", ""),
                                "Year": c.get("year", ""),
                                "Favorable?": f"{icon} {fav}",
                                "Strength": c.get("strength", ""),
                            })
                        st.dataframe(pd.DataFrame(table_data), use_container_width=True, hide_index=True)

                        for ci, c in enumerate(cases):
                            fav = c.get("favorability", "?")
                            border_color = "#22c55e" if fav == "FAVORABLE" else "#ef4444" if fav == "UNFAVORABLE" else "#94a3b8"
                            st.markdown(
                                f'<div style="border-left:3px solid {border_color};padding:10px 16px;'
                                f'margin:8px 0;border-radius:0 8px 8px 0;'
                                f'background:rgba(148,163,184,0.06);">'
                                f'<div style="font-weight:700;font-size:14px;">{c.get("citation","Unknown")}</div>'
                                f'<div style="font-size:12px;opacity:0.7;margin-bottom:6px;">{c.get("court","")} | {c.get("year","")}</div>'
                                f'<div style="font-size:13px;margin-bottom:4px;"><b>Holding:</b> {c.get("holding","")}</div>'
                                f'<div style="font-size:13px;margin-bottom:4px;"><b>Relevance:</b> {c.get("relevant_facts","")}</div>'
                                f'<div style="font-size:13px;margin-bottom:4px;"><b>Strategic Use:</b> {c.get("strategic_use","")}</div>'
                                f'</div>',
                                unsafe_allow_html=True,
                            )
                            quotes = c.get("key_quotes", [])
                            if quotes:
                                for kq in quotes[:3]:
                                    st.markdown(f"> *\"{kq}\"*")

                    next_searches = analysis.get("recommended_next_searches", [])
                    if next_searches:
                        st.markdown("**Suggested Follow-Up Searches**")
                        for ns in next_searches:
                            st.markdown(f"- `{ns}`")

            if st.button("Clear All Results", key="_clear_lex"):
                st.session_state["lexis_research_results"] = []
                st.rerun()

    with tabs[-1]:  # CHEAT SHEET (always last in Research & Draft)
        st.subheader("Quick-Reference Cheat Sheet")
        st.caption("One-page courtroom reference -- designed for the podium")

        if st.button("Generate Cheat Sheet", type="primary", key="_gen_cheat_sheet", help="Creates a condensed reference from your existing analysis (1 AI call)"):
            if results.get("case_summary"):
                with st.spinner("Generating courtroom cheat sheet..."):
                    cs_result = generate_cheat_sheet(st.session_state.agent_results)
                    st.session_state.agent_results["cheat_sheet"] = cs_result.get("cheat_sheet", "")
                    case_mgr.save_prep_state(case_id, prep_id, st.session_state.agent_results)
                    st.rerun()
            else:
                st.warning("Run a full analysis first so there's data to summarize.")

        cheat_sheet = results.get("cheat_sheet", "")
        if cheat_sheet:
            with st.container(border=True):
                st.markdown(cheat_sheet)

            st.divider()
            cs_col1, cs_col2 = st.columns(2)
            with cs_col1:
                st.download_button(
                    "Download as Text",
                    data=cheat_sheet,
                    file_name=f"cheat_sheet_{case_id}.txt",
                    mime="text/plain",
                    key="_dl_cheat_txt"
                )
            with cs_col2:
                try:
                    from core.export import generate_pdf_report
                    cs_pdf_state = {"case_summary": cheat_sheet, "strategy_notes": ""}
                    cs_pdf = generate_pdf_report(cs_pdf_state, f"Cheat Sheet -- {case_id}")
                    st.download_button(
                        "Download as PDF",
                        data=cs_pdf,
                        file_name=f"cheat_sheet_{case_id}.pdf",
                        mime="application/pdf",
                        key="_dl_cheat_pdf"
                    )
                except Exception:
                    pass
        else:
            st.info("No cheat sheet generated yet. Click the button above to create one from your analysis data.")

    # DEMAND LETTER GENERATOR
    _dl_tab_items = nav_groups.get("Research & Draft", nav_groups.get("\ud83d\udcda Research & Draft", []))
    _dl_tab_idx = -99
    for _i, _label in enumerate(_dl_tab_items):
        if "Demand" in _label:
            _dl_tab_idx = _i
            break

    if _dl_tab_idx != -99:
        with tabs[_dl_tab_idx]:
            st.subheader("Demand Letter Generator")
            st.caption("Generate professional demand letters for civil litigation with configurable tone, amounts, and deadlines.")

            demand_data = results.get("demand_letter", {})

            with st.expander("Letter Configuration", expanded=not bool(demand_data.get("letter_text"))):
                with st.form("demand_letter_form"):
                    st.markdown("##### Recipient")
                    dl_c1, dl_c2 = st.columns(2)
                    with dl_c1:
                        dl_recipient = st.text_input("Recipient Name", value="", placeholder="e.g., State Farm Insurance, Adjuster Jane Doe", key="_dl_recipient")
                    with dl_c2:
                        dl_recipient_addr = st.text_input("Recipient Address", value="", placeholder="123 Main St, Nashville, TN 37203", key="_dl_recipient_addr")

                    st.markdown("##### Demand")
                    dl_d1, dl_d2, dl_d3 = st.columns(3)
                    with dl_d1:
                        dl_amount = st.text_input("Demand Amount ($)", value="", placeholder="e.g., 250,000", key="_dl_amount")
                    with dl_d2:
                        dl_deadline = st.selectbox("Response Deadline", ["15 days", "30 days", "45 days", "60 days", "90 days"], index=1, key="_dl_deadline")
                    with dl_d3:
                        dl_tone = st.selectbox("Letter Tone", ["aggressive", "moderate", "conservative"], index=1, key="_dl_tone",
                            help="Aggressive = trial-ready posture. Moderate = firm but professional. Conservative = measured, good-faith.")

                    st.markdown("##### Attorney")
                    dl_a1, dl_a2 = st.columns(2)
                    with dl_a1:
                        dl_attorney = st.text_input("Attorney Name", value="", key="_dl_attorney")
                    with dl_a2:
                        dl_firm = st.text_input("Firm Name", value="", key="_dl_firm")

                    dl_instructions = st.text_area("Custom Instructions (Optional)", placeholder="e.g., Emphasize the defendant ran a red light, focus on future medical costs...", key="_dl_instructions", height=80)

                    dl_submit = st.form_submit_button("Generate Demand Letter", type="primary")

                if dl_submit:
                    if results.get("case_summary"):
                        with st.spinner("Generating demand letter -- analyzing case data, damages, and medical records..."):
                            params = {
                                "recipient": dl_recipient or "[Insurance Company / Opposing Party]",
                                "recipient_address": dl_recipient_addr or "[Address]",
                                "tone": dl_tone,
                                "demand_amount": dl_amount,
                                "deadline": dl_deadline,
                                "custom_instructions": dl_instructions,
                                "attorney_name": dl_attorney or "[Attorney Name]",
                                "firm_name": dl_firm or "[Firm Name]",
                            }
                            updates = generate_demand_letter(st.session_state.agent_results, params=params)
                            safe_update_and_save(case_mgr, st.session_state.current_case_id, st.session_state.current_prep_id, st.session_state.agent_results, updates)
                            st.rerun()
                    else:
                        st.warning("Analyze case documents first before generating a demand letter.")

            if demand_data and not demand_data.get("error"):
                letter_text = demand_data.get("letter_text", "")
                summary = demand_data.get("summary", "")
                sections = demand_data.get("sections", {})
                metadata = demand_data.get("metadata", {})

                if letter_text:
                    meta_cols = st.columns(4)
                    with meta_cols[0]:
                        st.metric("Tone", f"{metadata.get('tone', 'N/A').title()}")
                    with meta_cols[1]:
                        st.metric("Demand Amount", sections.get("demand_amount_stated", metadata.get("demand_amount", "N/A")))
                    with meta_cols[2]:
                        st.metric("Deadline", metadata.get("deadline", "N/A"))
                    with meta_cols[3]:
                        st.metric("Recipient", metadata.get("recipient", "N/A")[:30])

                    st.divider()

                    if summary:
                        st.info(f"**Summary**: {summary}")

                    st.markdown("### Full Demand Letter")
                    st.text_area("Letter Content", letter_text, height=500, key="_dl_full_text")

                    dl_col1, dl_col2, dl_col3 = st.columns(3)
                    with dl_col1:
                        st.download_button(
                            "Download as Text",
                            data=letter_text,
                            file_name=f"demand_letter_{metadata.get('date_generated', 'draft')}.txt",
                            mime="text/plain",
                            key="_dl_download_txt"
                        )
                    with dl_col2:
                        try:
                            from docx import Document as _DL_Doc
                            _dl_doc = _DL_Doc()
                            _dl_doc.add_heading("Demand Letter", 0)
                            for _para in letter_text.split("\n"):
                                if _para.strip():
                                    _dl_doc.add_paragraph(_para)
                            _dl_buf = io.BytesIO()
                            _dl_doc.save(_dl_buf)
                            _dl_buf.seek(0)
                            st.download_button(
                                "Download as Word",
                                data=_dl_buf,
                                file_name=f"demand_letter_{metadata.get('date_generated', 'draft')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="_dl_download_docx"
                            )
                        except Exception:
                            pass
                    with dl_col3:
                        try:
                            from core.export import generate_pdf_report
                            _dl_pdf_state = {"case_summary": letter_text, "strategy_notes": ""}
                            _dl_pdf = generate_pdf_report(_dl_pdf_state, f"Demand Letter -- {metadata.get('recipient', 'Draft')}")
                            st.download_button(
                                "Download as PDF",
                                data=_dl_pdf,
                                file_name=f"demand_letter_{metadata.get('date_generated', 'draft')}.pdf",
                                mime="application/pdf",
                                key="_dl_download_pdf"
                            )
                        except Exception:
                            pass

                    if sections:
                        st.divider()
                        st.markdown("### Section Breakdown")
                        sec_tabs = st.tabs(["Liability", "Injuries", "Damages"])
                        with sec_tabs[0]:
                            st.markdown(sections.get("liability", "_No liability section available._"))
                        with sec_tabs[1]:
                            st.markdown(sections.get("injuries", "_No injuries section available._"))
                        with sec_tabs[2]:
                            st.markdown(sections.get("damages_breakdown", "_No damages breakdown available._"))

            elif demand_data and demand_data.get("error"):
                st.error(f"Error generating demand letter: {demand_data['error']}")
            else:
                st.info("Configure the options above and click 'Generate Demand Letter' to create a professional demand letter.")
